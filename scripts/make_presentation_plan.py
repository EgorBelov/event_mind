"""Генератор docx с планом показа EventMind преподавателю."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "План_показа_EventMind.docx"


def set_base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def h1(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)


def h2(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2C, 0x55, 0x82)


def h3(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x3B, 0x6E, 0xA8)


def para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def script_block(doc: Document, what_show: str, what_say: str) -> None:
    """Блок «что показать / что говорить» — для каждого шага демо."""
    p = doc.add_paragraph()
    r = p.add_run("Что показать: ")
    r.bold = True
    p.add_run(what_show)

    p = doc.add_paragraph()
    r = p.add_run("Что говорить: ")
    r.bold = True
    p.add_run(what_say)


def build() -> None:
    doc = Document()
    set_base_style(doc)

    # ---------- Титульный блок ----------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("План показа курсовой работы")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("EventMind — агрегация IT-событий и персонализированные AI-рекомендации")
    r.italic = True
    r.font.size = Pt(13)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Прототип защиты перед преподавателем · ориентир: 15–20 минут")

    doc.add_paragraph()

    # ---------- 0. Подготовка ----------
    h1(doc, "0. Подготовка к защите (за 30 минут до показа)")

    para(doc, "Цель — чтобы во время демо ничего не пришлось ждать и устанавливать. Прогоните чек-лист заранее.")

    h3(doc, "Чек-лист окружения")
    bullet(doc, "Активировать venv: source .venv/bin/activate")
    bullet(doc, "Проверить .env: BOT_TOKEN, GROQ_API_KEY заполнены, DATABASE_URL=sqlite:///./eventmind.db")
    bullet(doc, "Накатить миграции: alembic upgrade head")
    bullet(doc, "В БД должны быть события: при необходимости curl -X POST http://localhost:8000/events/load")
    bullet(doc, "Открыть три терминала и заранее набрать команды (не запускать) — API, бот, scheduler")
    bullet(doc, "Открыть в браузере вкладки: http://localhost:8000/docs, http://localhost:8000/admin/, http://localhost:8000/health")
    bullet(doc, "Открыть Telegram-клиент с тестовым аккаунтом, найти бота, очистить чат через /start заново")
    bullet(doc, "Открыть в IDE/редакторе ключевые файлы: app/recommender/hybrid.py, app/agents/copilot/agent.py, app/api/main.py")
    bullet(doc, "Открыть README.md — это «шпаргалка», к ней можно отсылать преподавателя")
    bullet(doc, "Закрыть лишние окна и уведомления (на демо это сильно отвлекает)")

    h3(doc, "Резервный план (если что-то отвалится)")
    bullet(doc, "Если Groq упал по rate-limit — у вас работает fallback-модель, показать это как «фичу», а не «баг»")
    bullet(doc, "Если интернет лёг — есть offline-режим: rule-based рекомендации без LLM (показать /recommend)")
    bullet(doc, "Если бот не отвечает — показать те же сценарии через http://localhost:8000/docs (Swagger UI)")
    bullet(doc, "Иметь под рукой docs/Обзор_проекта_EventMind.docx и docs/отчет_Курсовая.docx — на случай вопросов по теории")

    doc.add_page_break()

    # ---------- 1. Структура выступления ----------
    h1(doc, "1. Структура выступления (тайминг)")

    timing = [
        ("1 мин", "Введение: о чём проект и какую задачу решает"),
        ("2 мин", "Архитектура и стек (на одном слайде / на диаграмме из README)"),
        ("2 мин", "Демо ingestion: как сырые события превращаются в нормализованные"),
        ("3 мин", "Демо Telegram-бота: профиль → /recommend → лайки/дизлайки/save"),
        ("2 мин", "AI-рекомендации: hybrid score + LLM-объяснения"),
        ("3 мин", "AI-Copilot (Supervisor-Worker, 5 специалистов) — главная фича"),
        ("1 мин", "Бэк-офис: /admin, /analytics, /trending"),
        ("1 мин", "Качество: тесты (167 шт.), миграции, fallback'и, /health"),
        ("1 мин", "Итоги: что сделано, какие технологии освоены, что дальше"),
        ("2–5 мин", "Ответы на вопросы"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Время"
    hdr[1].text = "Этап"
    for t, what in timing:
        row = table.add_row().cells
        row[0].text = t
        row[1].text = what

    doc.add_paragraph()
    para(doc, "Главное правило: 70 % времени — демо вживую, 30 % — рассказ. Преподаватель должен видеть результат на экране, а не только слайды.", bold=True)

    doc.add_page_break()

    # ---------- 2. Введение ----------
    h1(doc, "2. Введение (1 минута)")
    h3(doc, "Что говорить (текстом, можно зачитать)")
    para(
        doc,
        "«EventMind — это система, которая помогает IT-специалисту не утонуть в потоке "
        "митапов, конференций и онлайн-мероприятий. Я собираю события из публичных "
        "источников (Habr, RSS, KudaGo, Lu.ma, Telegram-каналы), нормализую их через "
        "LLM, строю профиль пользователя в Telegram и выдаю персональные рекомендации. "
        "Поверх этого работает AI-копилот: пользователь может задать цель — например, "
        "‘хочу за полгода вырасти в Data Engineer’ — и получить персональный roadmap "
        "из реальных событий».",
    )
    h3(doc, "Какую задачу решает")
    bullet(doc, "Проблема: событий много, но фильтры в источниках слабые (по теме «ML» приходит и для джунов, и митап про MLOps на Kubernetes)")
    bullet(doc, "Решение: профиль + история взаимодействий + embedding-сходство + Bayesian-предпочтения тем")
    bullet(doc, "Уникальность: не просто фильтр, а roadmap-планировщик под цель")

    # ---------- 3. Архитектура ----------
    h1(doc, "3. Архитектура и стек (2 минуты)")

    script_block(
        doc,
        what_show="Открыть README.md, найти раздел «Архитектура» с ASCII-диаграммой (около строки 80). Или открыть docs/diagrams/ если там есть картинки.",
        what_say="«Система состоит из трёх независимых процессов, чтобы их можно было масштабировать отдельно: FastAPI-бэкенд, Telegram-бот на aiogram, APScheduler для дайджестов и автоматического ingestion. Между ними — общая БД через SQLAlchemy 2.0».",
    )

    h3(doc, "Ключевые точки рассказа")
    bullet(doc, "Три процесса: API (uvicorn) + Bot (long-polling) + Scheduler (cron-job)")
    bullet(doc, "БД: SQLite в dev, PostgreSQL + pgvector в prod — переключается одним DATABASE_URL")
    bullet(doc, "LLM: Groq Cloud (llama-3.3-70b-versatile), своя обёртка _GroqWithFallback с автоматическим переключением на 8b при rate-limit'е")
    bullet(doc, "Оркестрация агентов: LangGraph (StateGraph) — Supervisor-Worker для Copilot'а")
    bullet(doc, "Эмбеддинги: sentence-transformers (multilingual MiniLM, 384-dim)")
    bullet(doc, "9 Alembic-миграций, все накатываются автоматически")

    h3(doc, "Если спросят «почему именно эти технологии»")
    bullet(doc, "FastAPI — async из коробки, автогенерация OpenAPI, типизация через pydantic")
    bullet(doc, "aiogram 3 — современный async-фреймворк под Telegram Bot API")
    bullet(doc, "LangGraph вместо LangChain Agents — нужен граф с условными переходами и состоянием")
    bullet(doc, "Groq — самый быстрый inference для open-source моделей, бесплатный tier для прототипа")
    bullet(doc, "SQLite вместо Postgres в dev — нулевая настройка для проверяющего")

    doc.add_page_break()

    # ---------- 4. Demo: ingestion ----------
    h1(doc, "4. Демо: пайплайн ingestion (2 минуты)")

    script_block(
        doc,
        what_show="Терминал 1: uvicorn app.api.main:app. Терминал 2: curl -X POST 'http://localhost:8000/ingestion/load-habr?limit=10' и затем curl http://localhost:8000/ingestion/status",
        what_say="«Источники у меня унифицированы под один интерфейс. Сырое событие сохраняется в таблицу raw_events, потом LangGraph-агент с Groq LLM достаёт из текста структурированный JSON: формат, город, темы, tech_stack, seniority, оценку качества и хайпа. Параллельно проверяется семантический дубликат по cosine ≥ 0.92 — иначе у меня появлялись копии одного митапа из Habr и RSS».",
    )

    h3(doc, "Что важно подсветить")
    bullet(doc, "Шесть источников: Habr (BS4), RSS, KudaGo, Lu.ma (ICS), Meetup (GraphQL), Telegram-каналы (web-preview scraping без api_id/api_hash)")
    bullet(doc, "LLM-нормализация с pydantic-валидацией — не доверяем «сырому» JSON от модели")
    bullet(doc, "Семантическая дедупликация (recommender/dedup.py) — ловит парафразы из разных источников")
    bullet(doc, "Не-IT события автоматически отсеиваются (topics == [] → status='non_it')")
    bullet(doc, "Эмбеддинг события считается сразу и кэшируется в events.embedding — рекомендации работают за ~0.4 c, а не 16 c")

    h3(doc, "Файлы, на которые можно ткнуть")
    bullet(doc, "app/ingestion/sources/ — все парсеры")
    bullet(doc, "app/agents/event_normalization/agent.py — LLM-агент нормализации")
    bullet(doc, "app/recommender/dedup.py — семантическая дедупликация")

    # ---------- 5. Demo: Telegram-бот ----------
    h1(doc, "5. Демо: Telegram-бот, профиль и рекомендации (3 минуты)")

    h3(doc, "Сценарий 5.1 — Onboarding и профиль")
    script_block(
        doc,
        what_show="В Telegram: /start → выбрать темы (AI/ML, Backend) → формат (online) → город. Затем /profile.",
        what_say="«Onboarding — четыре экрана с inline-кнопками. Профиль хранится в БД как topic_weights JSON, плюс параллельно по каждой теме держим Beta-распределение для Bayesian-предпочтений».",
    )

    h3(doc, "Сценарий 5.2 — Hybrid-рекомендации")
    script_block(
        doc,
        what_show="/recommend → показать карточку с like / dislike / save / похожие / следующее. Обратить внимание на строку «💡 Почему» с топ-3 компонентами score.",
        what_say="«Это hybrid-рекомендер. Каждое событие получает score из девяти компонентов: rule, cosine на embedding'ах, Bayesian Thompson, LinUCB bandit, LightGCN, quality, hype, freshness, skill-gap. Можно показать декомпозицию прямо в карточке».",
    )

    h3(doc, "Сценарий 5.3 — Обратная связь и адаптация")
    script_block(
        doc,
        what_show="Нажать 👍 на одной карточке, 👎 на другой. Затем /stats — показать, как обновились топ-темы. Потом /undo — откатить последний лайк.",
        what_say="«Каждый feedback обновляет четыре подсистемы одновременно: topic_weights, Bayesian α/β, LinUCB матрицу A и вектор b, плюс session-embedding последних пяти взаимодействий. Кнопка /undo делает rollback во всех четырёх — это нетривиально, обычно бандиты не умеют откатываться».",
    )

    h3(doc, "Сценарий 5.4 — Семантический поиск и cold-start")
    script_block(
        doc,
        what_show="/semantic нейросети для нлп — показать, что находит не только по точным словам. Потом /bio «Я backend-разработчик, хочу перейти в ML, опыт 3 года».",
        what_say="«Cold-start через bio: LLM с pydantic-валидацией извлекает темы, цели, текущие/целевые скиллы — заполняет UserSkillProfile и сразу пишет ‘виртуальные’ Bayesian-обновления, чтобы Thompson sampling не выдавал случайщину при первых запросах».",
    )

    doc.add_page_break()

    # ---------- 6. AI-рекомендации ----------
    h1(doc, "6. AI-рекомендации: hybrid + LLM-объяснения (2 минуты)")

    script_block(
        doc,
        what_show="В боте: «🎯 Рекомендации» → «AI». Показать, что объяснения — связный текст, а не шаблон. Засечь время — должно быть ~3 c, а не 15.",
        what_say="«Раньше у меня был LangGraph-граф из 4 нод, который делал по LLM-вызову на каждое событие — это занимало 15 секунд. Я переписал на ‘вариант C’: hybrid детерминированно за миллисекунды отбирает top-N кандидатов, потом один LLM-вызов с pydantic-схемой AIReasonBatch пишет объяснения батчем. Параметр ?limit=N реально работает от 1 до 10».",
    )

    h3(doc, "Что показать в коде, если попросят")
    bullet(doc, "app/api/services/ai_recommendation_service.py — единый батч-вызов")
    bullet(doc, "app/recommender/hybrid.py — score_breakdown с настраиваемыми весами")
    bullet(doc, "app/recommender/explain.py — explain_event_detailed + counterfactual («без skill-gap-сигнала это событие потеряло бы ~45 % веса»)")

    # ---------- 7. AI-Copilot ----------
    h1(doc, "7. AI-Copilot: Supervisor-Worker (3 минуты) — главная фича")

    para(
        doc,
        "Это самая нетривиальная часть проекта. На неё стоит потратить больше времени, чем на остальное демо.",
        bold=True,
    )

    h3(doc, "Сценарий 7.1 — Roadmap")
    script_block(
        doc,
        what_show="В боте: /copilot хочу за 6 месяцев вырасти из джуна в middle backend, нужен план мероприятий. Подождать ответ.",
        what_say="«LangGraph-граф: retrieve (RAG-выборка top-k событий + история + skill-профиль) → supervisor (классифицирует intent: recommend / career / roadmap / explain / summary) → один из пяти специалистов → finalize. Каждый специалист имеет свой system-prompt и подмножество function-calling tools».",
    )

    h3(doc, "Сценарий 7.2 — Career coach")
    script_block(
        doc,
        what_show="/copilot какие скиллы мне ещё нужно подтянуть для перехода в ML.",
        what_say="«Здесь сработал career-coach-специалист — у него фокус на skill_gap. Он сравнил target_skills из моего UserSkillProfile с tech_stack доступных событий и предложил конкретные пробелы».",
    )

    h3(doc, "Сценарий 7.3 — Объяснение события")
    script_block(
        doc,
        what_show="/copilot объясни почему ты порекомендовал событие #<id>.",
        what_say="«Сработал event-explainer: использует tool explain_event_detailed, разворачивает структурированный JSON-объяснение в человекочитаемый нарратив. Он не рекомендует новые события — только объясняет существующие. Это удобно для UX: пользователь не получит ‘спам’ новыми рекомендациями, когда задал уточняющий вопрос».",
    )

    h3(doc, "Что подсветить про архитектуру Copilot'а")
    bullet(doc, "Long-term memory (mem0-style): таблица user_memories — структурированные заметки о вкусах и целях, активно извлекаемые LLM из feedback'а и диалогов")
    bullet(doc, "Multi-turn: copilot_sessions хранит историю — Copilot помнит контекст прошлых сообщений")
    bullet(doc, "Supervisor — pydantic-валидированный JSON + keyword-fallback при сбое LLM (важно для надёжности)")
    bullet(doc, "Critique-revise loop в recommendation-специалисте — второй LLM-вызов критикует первый и исправляет")
    bullet(doc, "Function calling: 5 tools (search_events, get_user_profile, get_interactions_summary, save_for_user, explain_event)")

    h3(doc, "Файлы для показа в IDE")
    bullet(doc, "app/agents/copilot/agent.py — главный граф")
    bullet(doc, "app/agents/copilot/supervisor.py — intent-классификация")
    bullet(doc, "app/agents/copilot/specialists/ — все пять специалистов")
    bullet(doc, "app/agents/copilot/tools.py — function-calling")
    bullet(doc, "app/recommender/memory.py — long-term memory agent")

    doc.add_page_break()

    # ---------- 8. Бэк-офис ----------
    h1(doc, "8. Бэк-офис: admin-дашборд и аналитика (1 минута)")
    script_block(
        doc,
        what_show="Открыть http://localhost:8000/admin/ — таблица событий, пользователей, raw_events. Потом /trending в боте — ASCII-bar-график тем за 7 дней.",
        what_say="«HTML admin-дашборд — для оператора, чтобы видеть состояние БД. /analytics/* отдаёт JSON-агрегаты для топа тем и взаимодействий. /trending — пример агрегата прямо в боте».",
    )

    # ---------- 9. Качество кода ----------
    h1(doc, "9. Качество: тесты, миграции, надёжность (1 минута)")

    h3(doc, "Что показать")
    bullet(doc, "Запустить pytest -q — должны зелёные пройти 167 тестов (если не успеть — показать сам файл pyproject.toml с конфигом ruff/mypy/pytest)")
    bullet(doc, "Открыть alembic/versions/ — 9 миграций, включая pgvector для PostgreSQL")
    bullet(doc, "curl http://localhost:8000/health — композитный healthcheck (БД + прогрев эмбеддингов + Groq)")

    h3(doc, "Что говорить")
    para(
        doc,
        "«У меня 167 unit-тестов, покрывающих скоринг, Bayesian-обновления, бандит, GNN, всех специалистов Copilot'а, long-term memory и ingestion. "
        "Все внешние интеграции (Groq, источники) замоканы, тесты бегут оффлайн. Из инфраструктурных штук — Alembic-миграции, "
        "композитный /health, fallback Groq-модели и docker-compose из трёх сервисов с healthcheck'ом».",
    )

    # ---------- 10. Итоги ----------
    h1(doc, "10. Итоги (1 минута)")

    h3(doc, "Что сделано")
    bullet(doc, "Полноценный продукт из трёх процессов: API + бот + scheduler")
    bullet(doc, "6 источников ingestion с автоматической LLM-нормализацией и семантической дедупликацией")
    bullet(doc, "Hybrid-рекомендер из 9 компонентов с offline-метриками (Recall@k / nDCG@k)")
    bullet(doc, "AI-Copilot Supervisor-Worker с long-term memory и 5 специалистами")
    bullet(doc, "PostgreSQL + pgvector для прода, SQLite для dev")
    bullet(doc, "167 unit-тестов, 9 Alembic-миграций, Docker-compose")

    h3(doc, "Какие технологии освоены")
    bullet(doc, "LangGraph: StateGraph, conditional edges, structured output, function calling")
    bullet(doc, "Contextual bandits (LinUCB) и Bayesian (Thompson sampling) — на чистом numpy")
    bullet(doc, "LightGCN-style collaborative embeddings без torch")
    bullet(doc, "RAG-pipeline: sentence-transformers + cosine + MMR-диверсификация")
    bullet(doc, "pgvector + IVFFlat для прод-поиска по эмбеддингам")

    h3(doc, "Что можно улучшить (честно сказать — это произведёт лучшее впечатление)")
    bullet(doc, "Перенос на Postgres по умолчанию — сейчас prod-вариант, но dev на SQLite")
    bullet(doc, "GNN малополезен на маленьком датасете — выключен по умолчанию, нужен реальный поток данных")
    bullet(doc, "Нет авторизации в admin-дашборде — только локальный запуск")
    bullet(doc, "Frontend ограничен Telegram'ом — нет web-UI")

    doc.add_page_break()

    # ---------- 11. Ожидаемые вопросы ----------
    h1(doc, "11. Ожидаемые вопросы и заготовленные ответы")

    qa = [
        (
            "Почему hybrid, а не одна модель end-to-end?",
            "На прототипе с малым датасетом end-to-end-модель просто не на чем обучить. "
            "Hybrid из 9 компонентов даёт работающий baseline и интерпретируемость (видно, "
            "какой компонент сколько весит — есть score_breakdown и counterfactual).",
        ),
        (
            "Как вы боретесь с галлюцинациями LLM?",
            "Три уровня. Первый — pydantic-валидация всех JSON-ответов; невалидный JSON → "
            "keyword-fallback. Второй — LLM выбирает события из списка id, который я ему "
            "подал в промпт; парсер выкидывает id, которых нет в БД. Третий — RAG: LLM получает "
            "только релевантные события, а не весь каталог.",
        ),
        (
            "Чем LangGraph лучше LangChain Agents?",
            "Мне нужен явный граф с состоянием и условными переходами (intent → один из 5 "
            "специалистов). В LangChain Agents состояние неявное и трудно отлаживать. "
            "В LangGraph я вижу путь выполнения и могу подменить любую ноду.",
        ),
        (
            "Что будет, если упадёт Groq?",
            "Сначала _GroqWithFallback переключается на резервную модель (llama-3.1-8b-instant). "
            "Если упадёт и она — рекомендации отдаются с rule-based-объяснениями (полный фолбэк "
            "без даунтайма). Copilot вернёт ошибку только если упало всё.",
        ),
        (
            "Как тестируете LLM-агентов, если ответы не детерминированные?",
            "Моки. В тестах подменяю ChatGroq на stub, который возвращает фиксированный JSON. "
            "Проверяю не качество ответа модели, а правильность роутинга, парсинга, fallback'ов "
            "и обновления состояния. Качество ответа отдельно меряю через scripts/llm_judge.py.",
        ),
        (
            "Почему SQLite, а не Postgres?",
            "В dev — для нулевой настройки у проверяющего. В prod код переключается на "
            "Postgres + pgvector одним только DATABASE_URL — есть отдельная миграция "
            "9a1b2c3d4e5f, которая создаёт vector(384) + IVFFlat-индекс.",
        ),
        (
            "Что нового в подходе по сравнению с обычным collaborative filtering?",
            "Я использую гибрид: collaborative (LightGCN) + content-based (embedding'и описаний) + "
            "rule-based (явные предпочтения) + bandits (online-обучение). Главное — Copilot, "
            "который вообще не классическая задача рекомендаций: пользователь задаёт цель, "
            "а LLM строит persona-specific roadmap из реальных событий.",
        ),
        (
            "Сколько строк кода и сколько времени заняло?",
            "Около 15 тысяч строк Python в app/, плюс 167 тестов. "
            "Конкретное время — назвать честно. Если не помните точно, скажите ‘примерно три месяца "
            "активной разработки’ — это нормальный диапазон для курсовой такого объёма.",
        ),
        (
            "Какие метрики качества рекомендаций?",
            "Offline — leave-one-out Recall@5 и nDCG@5 в scripts/eval_offline.py, "
            "три варианта скорера (rule / hybrid / bayesian). Плюс LLM-as-judge "
            "(scripts/llm_judge.py) — оценивает релевантность и diversity-contribution по 1–5.",
        ),
        (
            "Можно ли это запустить у себя?",
            "Да: git clone, pip install -r requirements.txt, alembic upgrade head, "
            "положить GROQ_API_KEY и BOT_TOKEN в .env, uvicorn app.api.main:app. "
            "Либо docker compose up -d --build — поднимает все три сервиса сразу. "
            "Всё описано в README.md.",
        ),
    ]
    for q, a in qa:
        p = doc.add_paragraph()
        r = p.add_run("В: ")
        r.bold = True
        p.add_run(q)
        p = doc.add_paragraph()
        r = p.add_run("О: ")
        r.bold = True
        p.add_run(a)
        doc.add_paragraph()

    # ---------- 12. Шпаргалка ----------
    h1(doc, "12. Шпаргалка: команды для копи-паста")

    h3(doc, "Запуск всего")
    para(doc, "Терминал 1 (API):")
    para(doc, "  source .venv/bin/activate && uvicorn app.api.main:app --reload")
    para(doc, "Терминал 2 (бот):")
    para(doc, "  source .venv/bin/activate && python -m app.bot.main")
    para(doc, "Терминал 3 (scheduler, опционально):")
    para(doc, "  source .venv/bin/activate && python -m app.scheduler.digest")

    h3(doc, "Ingestion на демо")
    para(doc, "  curl -X POST http://localhost:8000/events/load")
    para(doc, "  curl -X POST 'http://localhost:8000/ingestion/load-habr?limit=10'")
    para(doc, "  curl    http://localhost:8000/ingestion/status")

    h3(doc, "Проверки")
    para(doc, "  curl http://localhost:8000/health")
    para(doc, "  open http://localhost:8000/docs")
    para(doc, "  open http://localhost:8000/admin/")

    h3(doc, "Тесты (если попросят запустить)")
    para(doc, "  pytest -q")
    para(doc, "  pytest tests/test_supervisor.py -v   # отдельно про Copilot")
    para(doc, "  pytest tests/test_bandit.py -v       # отдельно про LinUCB")

    h3(doc, "Сценарии в Telegram-боте (по порядку)")
    numbered(doc, "/start — пройти onboarding")
    numbered(doc, "/profile — показать собранный профиль")
    numbered(doc, "/recommend — лента hybrid-рекомендаций, нажать 👍 и 👎")
    numbered(doc, "Кнопка «AI-рекомендации» — показать LLM-объяснения батчем")
    numbered(doc, "/stats — увидеть, как обновились веса")
    numbered(doc, "/undo — откатить последний feedback")
    numbered(doc, "/semantic нейросети для нлп — показать семантический поиск")
    numbered(doc, "/bio Я backend, хочу в ML — cold-start через LLM")
    numbered(doc, "/copilot хочу за 6 месяцев вырасти в middle backend — roadmap")
    numbered(doc, "/copilot какие скиллы подтянуть — career coach")
    numbered(doc, "/trending — ASCII-bar-график тем за неделю")

    # ---------- 13. Финальные советы ----------
    h1(doc, "13. Финальные советы по защите")
    bullet(doc, "Не пытайтесь показать ВСЁ — выберите 3–4 фичи, которые «вау», и сфокусируйтесь на них")
    bullet(doc, "Главные «вау»-моменты: AI-Copilot с roadmap'ом, hybrid score_breakdown в карточке, /undo с откатом 4 подсистем, fallback-цепочка Groq")
    bullet(doc, "Если что-то падает — не паникуйте, скажите «вот тут как раз срабатывает fallback, смотрите» и переходите к плану B")
    bullet(doc, "Честно говорите про ограничения — преподаватели ценят рефлексию выше, чем показуху")
    bullet(doc, "Держите README.md открытым — там есть ответ почти на любой технический вопрос")
    bullet(doc, "Если вопрос вне области — говорите «это вне scope текущего прототипа, в следующей итерации сделаю так-то»")

    # Сохранить
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()

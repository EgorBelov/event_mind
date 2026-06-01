"""Полная регенерация трёх docs/-документов под текущую реализацию проекта.

Что переписывается:
  - docs/Обзор_проекта_EventMind.docx        — стиль H1/H2/H3 + Normal/List Bullet
  - docs/План_показа_EventMind.docx          — стиль Heading 1/2/3 + Normal/List Bullet
  - docs/отчет_Курсовая.docx                  — стиль HSE Title 1/2/3 + HSE Default Text + HSE list

Скрипт:
  1) Открывает существующий .docx (чтобы сохранить определения стилей и sectPr).
  2) Полностью вычищает body (кроме финального sectPr).
  3) Строит контент заново — из актуального состояния проекта.

Запуск:
    python -m scripts.regenerate_docs

Скрипт идемпотентен: повторный запуск даёт тот же результат.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document

DOCS = Path(__file__).resolve().parents[1] / "docs"

# Текущая дата — дата сборки документов.
BUILD_DATE = "29 мая 2026 г."


# ─── Очистка тела документа ──────────────────────────────────────────────


def _clear_body(doc) -> None:
    """Удалить все child-элементы body, кроме финального sectPr."""
    body = doc.element.body
    to_remove = [
        child for child in list(body)
        if child.tag.rsplit("}", 1)[-1] != "sectPr"
    ]
    for el in to_remove:
        body.remove(el)


# ─── Builder с привязкой к стилям конкретного документа ──────────────────


class Builder:
    def __init__(self, doc, *, h1: str, h2: str, h3: str, normal: str, bullet: str):
        self.doc = doc
        self._h1 = h1
        self._h2 = h2
        self._h3 = h3
        self._normal = normal
        self._bullet = bullet

    def h1(self, text: str) -> None:
        self.doc.add_paragraph(text, style=self._h1)

    def h2(self, text: str) -> None:
        self.doc.add_paragraph(text, style=self._h2)

    def h3(self, text: str) -> None:
        self.doc.add_paragraph(text, style=self._h3)

    def p(self, text: str) -> None:
        self.doc.add_paragraph(text, style=self._normal)

    def bullet(self, text: str) -> None:
        self.doc.add_paragraph(text, style=self._bullet)

    def empty(self) -> None:
        self.doc.add_paragraph("", style=self._normal)


# ════════════════════════════════════════════════════════════════════════
#  Обзор проекта EventMind
# ════════════════════════════════════════════════════════════════════════


def build_overview(b: Builder) -> None:
    b.h1("Обзор проекта EventMind")
    b.p(
        "Документ описывает актуальное состояние проекта EventMind по итогам "
        "майской итерации 2026. Здесь зафиксированы: какие подсистемы работают, "
        "какие компоненты алгоритма ранжирования включены, какие правки были "
        "сделаны за последнюю неделю и что остаётся открытым."
    )
    b.p(f"Дата сборки документа: {BUILD_DATE}")

    # ── 1. Описание ──────────────────────────────────────────────────────
    b.h2("1. Описание проекта")
    b.p(
        "EventMind — система агрегации IT-мероприятий и персонализированных "
        "рекомендаций. Архитектура построена из трёх независимых процессов: "
        "REST API на FastAPI, Telegram-бот на aiogram 3 и фоновый scheduler "
        "на APScheduler. Каждый процесс работает в своём контейнере и "
        "общается с другими через REST и общую базу данных."
    )
    b.p(
        "Рекомендательная подсистема — гибридная. На каждый запрос параллельно "
        "считаются девять компонентов: правила по совпадению профиля, "
        "семантическая близость по эмбеддингам, байесовская модель "
        "предпочтений тем (Thompson sampling), скоры качества и актуальности, "
        "свежесть даты, соответствие карьерным целям, контекстный бандит "
        "LinUCB и коллаборативный сигнал LightGCN. Итог взвешивается, "
        "сортируется и пропускается через MMR-диверсификацию."
    )
    b.p(
        "AI-слой реализован на LangGraph и langchain-groq. Endpoint "
        "/agent-recommendations берёт top-N от hybrid и одним вызовом LLM "
        "получает объяснения. Endpoint /copilot/{tid}/turn — мульти-туровый "
        "диалог: Supervisor-Worker граф с пятью специалистами и "
        "поддерживаемой через таблицу copilot_sessions памятью."
    )
    b.p(
        "Подсистема ingestion включает шесть адаптеров: Habr, RSS/Atom, "
        "KudaGo, Lu.ma (ICS), Meetup (GraphQL), Telegram-каналы (web-preview). "
        "Каждое сырое сообщение проходит через LLM-нормализацию с "
        "Pydantic-валидацией, дедупликацию по cosine ≥ 0.92 и записывается "
        "в events. Запуск — периодически через scheduler, разделение по "
        "интервалам через INGEST_INTERVAL_HOURS."
    )

    b.h3("1.1. Ключевой стек")
    b.bullet("Python 3.12, FastAPI 0.110, Uvicorn, aiogram 3.13")
    b.bullet("SQLAlchemy 2.0 + Alembic; SQLite (dev fallback) и PostgreSQL + pgvector на Supabase (prod/dev)")
    b.bullet(
        "LangGraph 0.2, langchain-core; multi-provider LLM-цепочка _LLMChain: "
        "Gemini (REST-транспорт, автопроба модели на старте) → Groq llama-3.3-70b → Groq llama-3.1-8b. "
        "Circuit-breaker + per-provider cooldown."
    )
    b.bullet("sentence-transformers 2.7 — мультиязычный MiniLM-L12-v2 (384 dim)")
    b.bullet("APScheduler 3.10 — daily_digest, ingest_*, backfill_event_embeddings, backfill_user_embeddings, compact_memories")
    b.bullet("feedparser, BeautifulSoup 4 + lxml, httpx, pydantic 2")
    b.bullet("pytest 8.2 — 285 тестов проходят, ruff — без ошибок; CI с отдельным test-postgres job на pgvector/pgvector:pg16")
    b.bullet("Docker + docker-compose: три контейнера, healthcheck на бэке перед стартом bot/scheduler")

    # ── 2. Что сделано ────────────────────────────────────────────────────
    b.h2("2. Что сделано")

    b.h3("2.1. Состав репозитория")
    b.bullet("app/api — REST: routers users, events, recommendations, agent_recommendations, copilot, ingestion, analytics, subscriptions, vocabulary, admin")
    b.bullet("app/bot — handlers start, recommendations, profile, search, copilot, subscriptions + reply/inline-клавиатуры")
    b.bullet("app/recommender — hybrid, scoring, embeddings, bayesian, bandit, gnn, skill_gap, retrieval, diversity, dedup, memory, explain")
    b.bullet("app/agents — recommendation (event-нормализация), copilot (supervisor + 5 специалистов + tools)")
    b.bullet("app/ingestion — sources/{habr,rss,kudago,luma,meetup,tg_channels}")
    b.bullet("app/scheduler — digest + ingestion + memory compaction на APScheduler")
    b.bullet("app/db/models — user, event, raw_event, interaction, topic, copilot_session, user_topic_stat, user_bandit_state, user_skill_profile, user_memory")
    b.bullet("alembic — 14 миграций; включая pgvector, copilot_sessions, bandit_state, user_memories, skill_profile, enriched_fields, retry_count + start_at, feed_cursor, recommendation_cache, series_slug, telegram_id BIGINT")
    b.bullet("tests — ~35 файлов, 285 кейсов: hot-path get_recommendations, recommendation_cache, feed_cursor, series, ingestion (idempotent + adaptive batching), enum-validation, dedup, supervisor, copilot tools, LLM circuit-breaker, Gemini probe, prompt safety, API key auth, request middleware, config validate, date localization, scoring, bayesian, bandit, gnn, memory hard-cap, cold-start, scheduler, pgvector")
    b.bullet("scripts — train_gnn, eval_offline (Recall@k/nDCG@k), llm_judge, compact_memories, backfill_pgvector, make_presentation_plan, append_overview_chapter")

    b.h3("2.2. Готовая функциональность")
    b.bullet(
        "REST API: пользователи, события, рекомендации, AI-рекомендации, Copilot "
        "(one-shot и multi-turn), ingestion (Habr, RSS, KudaGo, Lu.ma, Meetup, TG), "
        "subscriptions, аналитика трендов, словари тем/городов/форматов, admin-эндпоинты."
    )
    b.bullet(
        "Telegram-бот: тур по командам (4 экрана) с возможностью пропустить "
        "и сразу попасть на запуск настройки; мастер настройки профиля "
        "(темы → формат → город); reply-меню «🎯 Рекомендации / 🔍 Поиск / "
        "👤 Профиль / ⚙️ Ещё»; карточка события с кнопками ❓ Почему / "
        "📖 Подробнее / 👍 / 👎 / ⭐ / Похожие / Следующее."
    )
    b.bullet(
        "Hybrid recommender: 9 компонентов (rule, cosine, bayesian, quality, "
        "hype, freshness, skill_gap, bandit, gnn), каждая компонента изолирована "
        "try/except — сбой одной не ломает выдачу."
    )
    b.bullet(
        "MMR-диверсификация: λ=0.7 по умолчанию, переключается через "
        "MMR_ENABLED, применяется поверх отсортированного списка."
    )
    b.bullet(
        "AI-рекомендации (вариант C): hybrid отбирает top-N, один LLM-вызов "
        "с pydantic-схемой пишет объяснения батчем — быстрее старой 4-нодовой "
        "цепочки и без жёсткого лимита в ровно 3 карточки."
    )
    b.bullet(
        "AI Copilot: Supervisor-Worker граф LangGraph с пятью специалистами "
        "(recommendation, career_coach, roadmap, explainer, summary); шесть "
        "function-calling tools (search_events, get_user_profile, "
        "get_interactions_summary, explain_event, recall_about_user, mark_saved); "
        "история диалога — в copilot_sessions, multi-turn память; long-term "
        "memory подмешивается в recall."
    )
    b.bullet(
        "Долговременная память пользователя (стиль mem0): write/recall/extract/"
        "compact с фильтром salience и LLM-агрегацией заметок по категориям; "
        "compaction job раз в 7 дней."
    )
    b.bullet(
        "Ingestion: 6 адаптеров источников + AI-нормализация raw_events → "
        "events с Pydantic-валидацией; периодический запуск через APScheduler "
        "по INGEST_INTERVAL_HOURS; non-IT события отсеиваются."
    )
    b.bullet("Семантическая дедупликация — cosine ≥ 0.92 за 90 дней.")
    b.bullet(
        "Эвал: scripts/eval_offline.py (leave-one-out, Recall@k/nDCG@k для "
        "rule/hybrid/bandit) и scripts/llm_judge.py (LLM-судья поверх ленты)."
    )
    b.bullet(
        "Деплой: Dockerfile + docker-compose; три сервиса (api, bot, "
        "scheduler); healthcheck на бэке перед стартом зависимых; "
        "INGEST_ENABLED, GNN_ENABLED, MEMORY_ENABLED — фичефлаги."
    )

    # ── 3. Доработки мая 2026 ────────────────────────────────────────────
    b.h2("3. Доработки мая 2026")

    b.h3("3.0. Новейшее: Postgres/Supabase, масштабирование, ops")
    b.bullet(
        "Dev-БД переведена на управляемый PostgreSQL (Supabase, session pooler): "
        "pgvector активен, retrieval и кандидатный отбор идут через индекс <=>; "
        "engine с pool_pre_ping для устойчивости к простаивающим соединениям."
    )
    b.bullet(
        "Масштабирование /recommendations: кандидатный отбор top-N через pgvector "
        "(вместо скоринга всего каталога) + joinedload против N+1 на темах; "
        "тяжёлые объяснения строятся только для возвращаемого top-N."
    )
    b.bullet(
        "Ingestion: единый POST /ingestion/load-all (все источники по очереди), "
        "POST /ingestion/retry-failed, батчевая LLM-нормализация (несколько событий "
        "на один вызов — экономия токенов) с early-stop на rate-limit (429/TPD); "
        "embedding_vec пишется сразу при нормализации."
    )
    b.bullet(
        "Данные: events.start_at (DateTime) для freshness/сортировки + "
        "raw_events.retry_count (retry-failed пропускает безнадёжные после N попыток)."
    )
    b.bullet(
        "Ops/качество: scheduler на logging; ruff = 0 ошибок по репо; CI на GitHub "
        "Actions (ruff + pytest); LLM-обёртка строится лениво (импорт не требует ключа)."
    )

    b.h3("3.1. Свежий пакет правок (июнь 2026)")
    b.bullet(
        "GET /recommendations стал read-only: refresh_user_embedding и "
        "ensure_event_embeddings удалены из горячего пути (под Supabase "
        "каждый GET был writer, по UPDATE на каждое открытие ленты). Прогрев "
        "user.embedding перенесён на пишущие пути (register/edit/analyze-bio/"
        "feedback). Прогрев event.embedding — на ingestion + scheduler-джоб "
        "backfill_event_embeddings (раз в час) и backfill_user_embeddings (раз в 4 ч)."
    )
    b.bullet(
        "TTL-кэш рекомендаций: таблица recommendation_cache (telegram_id PK + "
        "items_json + expires_at). Hit пропускает весь скоринг. TTL по умолчанию "
        "15 минут; инвалидация на feedback / register / edit / analyze-bio."
    )
    b.bullet(
        "Multi-provider LLM-цепочка _LLMChain: Gemini первым (через REST-транспорт, "
        "т.к. gRPC-резолвер у Google на macOS/AdGuard режется по DNS), Groq 70b и 8b "
        "fallback'ами. Автопроба Gemini-модели на старте API: HTTP-POST по списку "
        "кандидатов, берём первую с 200 — кэшируется до перезапуска или "
        "POST /admin/llm/reprobe. Circuit-breaker (5 фейлов цепочки → cooldown 120 с) + "
        "per-provider cooldown (2 фейла одного звена → skip 10 мин)."
    )
    b.bullet(
        "Security-слой: ApiKeyMiddleware (X-API-Key через hmac.compare_digest) + "
        "RequestLogMiddleware (одна строка на запрос с тайминговым суффиксом и "
        "X-Request-ID) + exception_handler для HTTPException 4xx (WARNING без "
        "traceback вместо стандартного ERROR-trace на нормальные 404)."
    )
    b.bullet(
        "Prompt-injection sanitize: модуль app/core/prompt_safety.py "
        "(контроль-символы / zero-width / BIDI чистка + обёртка <user_input> + "
        "лог injection-фраз). Подключено в bio-обработку и copilot."
    )
    b.bullet(
        "Анти-флуд серий: events.series_slug (snake_case, снимает #15, vol.2, "
        "годы, римские в контексте) — в /recommendations оставляем один выпуск "
        "серии (ближайший к now по start_at)."
    )
    b.bullet(
        "Строгая enum-валидация полей нормализатора (format/event_type/seniority — "
        "closed-domain, city/level — open-domain с anti-malformed-фильтром) + "
        "ISO-date validation с календарной проверкой и диапазоном лет."
    )
    b.bullet(
        "Idempotent ingestion: при повторной нормализации того же события "
        "обновляются только поля, которые стали информативнее (была пустая дата → "
        "появилась; format=unknown → online; описание стало длиннее на 20%+). "
        "Раньше re-ingestion просто пропускал — терялись обновления анонсов."
    )
    b.bullet(
        "Курсор ленты бота перенесён из module-dict в users.feed_cursor "
        "(переживает рестарт, работает под несколькими воркерами). "
        "users.telegram_id переведён на BIGINT (каналы/группы имеют ID > 2^31)."
    )
    b.bullet(
        "Memory hard-cap = 500 заметок на пользователя; copilot-история сжимается "
        "одним LLM-summary turn'ом при превышении окна (вместо тупого среза)."
    )
    b.bullet(
        "Локализация дат в боте: events.start_at рендерится в карточке с учётом "
        "таймзоны (Europe/Moscow по умолчанию)."
    )

    b.h3("3.2. PostgreSQL + pgvector")
    b.p(
        "Миграция 9a1b2c3d4e5f добавляет тип vector(384) к events.embedding и "
        "users.embedding для PostgreSQL; на SQLite колонка остаётся TEXT "
        "(JSON-сериализация). Функция pgvector_top_k_events использует "
        "оператор <=> (cosine distance) для top-k вместо in-process сортировки. "
        "Скрипт scripts/backfill_pgvector.py заполняет существующие "
        "embedding-кэши после перехода на PG."
    )

    b.h3("3.3. Композитный /health и LLM-цепочка")
    b.p(
        "Endpoint /health возвращает {db: ok, embeddings: ok, llm: ok} — "
        "каждый компонент проверяется независимо. Поле llm показывает реально "
        "выбранную Gemini-модель (по автопробе) и сконфигурированные Groq-fallback'и. "
        "Цепочка _LLMChain в app/agents/recommendation/llm.py идёт по звеньям до "
        "первого успеха: Gemini → Groq 70b → Groq 8b. Контракт invoke / "
        "with_structured_output / bind_tools совместим со старым ChatGroq — "
        "пользователи llm.invoke ничего не меняют. Доступен POST /admin/llm/reprobe "
        "для повторного выбора Gemini-модели без рестарта (когда квота на текущей "
        "выбранной кончилась)."
    )

    b.h3("3.4. AI-рекомендации, вариант C")
    b.p(
        "Старая четырёхнодовая цепочка LangGraph (user_profile → "
        "event_analyzer → recommendation → explanation) убрана. Заменено на: "
        "hybrid отбирает top-N кандидатов детерминированно, один LLM-вызов "
        "с pydantic-схемой выписывает короткие объяснения для каждого. "
        "Параметр limit реально влияет на длину выдачи."
    )

    b.h3("3.5. UX Telegram-бота")
    b.p(
        "Введён мульти-туровый Copilot: новый файл handlers/copilot.py, "
        "точка входа /copilot или callback из поиска, активная сессия "
        "живёт 15 минут; продолжение диалога — простыми сообщениями. Карточка "
        "события переработана: вместо переключения mode=rule/ai теперь два "
        "независимых пояснения — ❓ Почему (короткий LLM-ответ в поп-апе) и "
        "📖 Подробнее (развёрнутый разбор + советы + метрики)."
    )

    b.h3("3.6. Качество кода")
    b.p(
        "ruff на app/ — без ошибок. Тестовый набор — 285 кейсов в ~35 файлах. "
        "Прямо покрыты hot-path get_recommendations (включая cache-hit пропускает "
        "scoring через spy), TTL-кэш, курсор ленты, series anti-flood, "
        "idempotent ingestion, enum-валидация, LLM-цепочка (5 тестов на chain + "
        "circuit-breaker + Gemini probe), prompt-injection sanitize, "
        "ApiKeyMiddleware (5 тестов), RequestLogMiddleware (3 теста), "
        "fail-fast config, локализация дат. Pytest-фикстуры используют "
        "in-memory SQLite, чтобы тесты не конфликтовали с dev-БД. "
        "CI с двумя job'ами: test-sqlite на каждом push/PR и test-postgres "
        "на сервисе pgvector/pgvector:pg16 — критичные backend-зависимые "
        "пути (dedup, retrieval, recommendation_cache, feed_cursor, "
        "ingestion, pgvector) гоняются на реальном Postgres."
    )

    # ── 4. Что осталось открытым ──────────────────────────────────────────
    b.h2("4. Что осталось открытым")
    b.bullet(
        "GNN отключён по умолчанию (gnn_enabled=False): на текущем размере "
        "графа (порядок десятков лайков) коллаборативный сигнал шумит. "
        "Включается через scripts/train_gnn.py + переменную окружения."
    )
    b.bullet(
        "Адаптер Meetup требует валидного OAuth Pro-токена — на демо "
        "используется пустой MEETUP_TOKEN и Meetup-источник пропускается."
    )
    b.bullet(
        "Lu.ma-адаптер работает по ICS-фидам конкретных календарей "
        "(LUMA_CALENDARS) — не публичный API."
    )
    b.bullet(
        "Long-term memory compaction срабатывает раз в неделю (по умолчанию). "
        "При интенсивных тестах можно дёргать вручную: python -m scripts.compact_memories."
    )
    b.bullet(
        "На dev-SQLite WAL + busy_timeout=10s, но при долгих параллельных "
        "ingestion-джобах единичные «database is locked» возможны. На PG "
        "проблема не воспроизводится."
    )

    # ── 5. Состояние документации ────────────────────────────────────────
    b.h2("5. Состояние документации")
    b.p(
        "Этот файл (Обзор_проекта_EventMind.docx) — авторитетный summary "
        "архитектуры и состояния. План защиты лежит в "
        "План_показа_EventMind.docx (главы 1–14, включая глоссарий и "
        "сценарии). Полный технический отчёт оформлен в "
        "отчет_Курсовая.docx по шаблону HSE: введение, анализ предметной "
        "области, проектирование, реализация, тестирование, доработки, "
        "заключение."
    )
    b.p(
        "В docs/diagrams лежат PNG-диаграммы: общая архитектура, поток "
        "ingestion, hybrid-конвейер, Supervisor-Worker, ER-диаграмма. Они "
        "ссылаются из всех трёх документов и из presentation-плана."
    )

    # ── 6. Итог ───────────────────────────────────────────────────────────
    b.h2("6. Итог")
    b.p(
        "EventMind — рабочий end-to-end прототип hybrid-рекомендера для IT-"
        "событий с AI-копилотом. На стороне рекомендера девять компонент "
        "одновременно реализуют content-based, контекстный (LinUCB), "
        "коллаборативный (LightGCN), Bayesian-preference и эвристические "
        "(quality, hype, freshness, skill_gap) подходы. На стороне UX — "
        "Telegram-бот с настройкой профиля, лентой рекомендаций, "
        "объяснимостью двух уровней (короткое и развёрнутое) и мульти-"
        "туровый AI-копилот с долговременной памятью. На стороне данных — "
        "шесть ingestion-источников с AI-нормализацией и семантической "
        "дедупликацией. Все ключевые компоненты покрыты тестами; "
        "критические пути ловят сбои LLM и БД через fallback и rollback "
        "и не разваливают пользовательский сценарий."
    )

    # ── 7. Рекомендательная модель — подробная глава ─────────────────────
    _build_recsys_chapter(b)


def _build_recsys_chapter(b: Builder) -> None:
    """Глава о рекомендательной модели — портирована из append_overview_chapter.py."""

    b.h1("7. Рекомендательная модель: устройство пользователя, события и 9 компонент")
    b.p(
        "Глава вставлена скриптом scripts/regenerate_docs.py и является "
        "авторитетным описанием алгоритма ранжирования. При расхождении "
        "с другими разделами — здесь источник истины."
    )

    # 7.1 Модель пользователя
    b.h2("7.1. Модель пользователя — какие слои её составляют")
    b.p(
        "Модель пользователя — это шесть взаимодополняющих слоёв. Каждый "
        "слой обновляется в своё время и используется разными компонентами "
        "hybrid-скоринга."
    )

    b.h3("Слой 1. Базовый профиль (таблица users)")
    b.p(
        "Хранит идентификаторы и явные предпочтения: telegram_id, username, "
        "preferred_format (online/offline/hybrid/any), city, флаг "
        "is_subscribed. Здесь же JSON-словарь topic_weights — веса интересов "
        "по темам (целое, по умолчанию 5). Обновляется при онбординге, /edit "
        "и каждом фидбэке."
    )
    b.p(
        "В этой же таблице кэшируется embedding — 384-мерный вектор от "
        "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2. "
        "Используется компонентом cosine; пересчитывается при изменении "
        "профиля или вызове GET /recommendations (с обязательным rollback "
        "на случай «database is locked»)."
    )

    b.h3("Слой 2. Темы-пользователя (таблица user_topics)")
    b.p(
        "Связь многие-ко-многим: какие темы пользователь выбрал явно. "
        "Используется в rule-компоненте. Обновляется только при изменении "
        "профиля через onboarding или /edit."
    )

    b.h3("Слой 3. Bayesian-статистика по темам (таблица user_topic_stats)")
    b.p(
        "На каждую пару (пользователь, тема) хранятся alpha и beta — "
        "параметры Beta-распределения. Лайк/save увеличивают alpha, дизлайк "
        "увеличивает beta. При ранжировании событие получает сигнал по "
        "Thompson sampling: для каждой темы события сэмплируется p из "
        "Beta(α, β) и берётся максимум — авто-баланс explore/exploit."
    )
    b.p(
        "Реализован temporal decay: при чтении α и β затухают к prior со "
        "скоростью γ^days_since_update (BAYES_DECAY_PER_DAY=0.995). Decay "
        "ленивый — в load_user_stats, отдельный job не нужен."
    )

    b.h3("Слой 4. Skill-профиль (таблица user_skill_profiles)")
    b.p(
        "Структурированное описание роли: current_role, target_role, "
        "seniority, JSON-массивы current_skills, target_skills, goals, "
        "learning_horizon. Заполняется через /bio (cold-start): LLM с "
        "pydantic-валидацией извлекает поля из свободного текста. "
        "Используется skill_gap-компонентом."
    )

    b.h3("Слой 5. Состояние контекстного бандита (таблица user_bandit_states)")
    b.p(
        "Per-user онлайн-обучаемая линейная регрессия LinUCB: dim=16, "
        "a_json (d×d матрица A), b_json (d-вектор b), update_count. "
        "Контекстный вектор x формируется в context_vector: топ-3 темы "
        "пользователя, формат, город, freshness и др."
    )
    b.p(
        "На каждый фидбэк: A += x·xᵀ, b += reward·x (reward = +1 / +1.5 / "
        "−1). Скор — UCB: xᵀθ̂ + α·√(xᵀA⁻¹x), где θ̂ = A⁻¹·b. Чистый numpy, "
        "без torch."
    )

    b.h3("Слой 6. Долговременная память (таблица user_memories)")
    b.p(
        "Активные заметки про пользователя от первого лица: «хочет перейти "
        "в ML за 6 месяцев», «активно сохраняет Kubernetes-события». "
        "Заметки пишутся LLM-агентом из значимого фидбэка и реплик "
        "Copilot'а. У каждой: category, salience (1–10), embedding, "
        "access_count, last_accessed_at, source."
    )
    b.p(
        "Используется специалистами Copilot через tool recall_about_user: "
        "top-k по score = 0.7·cosine + 0.3·salience_norm. При накоплении "
        ">50 заметок LLM группирует по категориям и сжимает до ≤5 на "
        "категорию (job раз в 7 дней)."
    )

    # 7.2 Модель события
    b.h2("7.2. Модель события — что и откуда")
    b.p(
        "Событие проходит две стадии: raw_events (сырой текст) → events "
        "(нормализованные поля). Все поля заполняет LLM-агент "
        "event_normalization при ingestion, результат валидируется "
        "Pydantic-схемой."
    )

    b.h3("Поля нормализованного события (events)")
    b.bullet("title, description — основной контент")
    b.bullet("format (online/offline/hybrid/any), city, level, date — фильтры")
    b.bullet("topics (через event_topics) — many-to-many с топиками")
    b.bullet("summary — краткое содержание для поиска и отображения")
    b.bullet("tech_stack (JSON) — список технологий: Kubernetes, MLflow, PostgreSQL и т.д.")
    b.bullet("seniority — целевая аудитория по опыту (используется skill_gap)")
    b.bullet("quality_score (1–10) — оценка содержательности от LLM")
    b.bullet("hype_score (1–10) — оценка хайповости от LLM")
    b.bullet(
        "embedding — 384-мерный вектор от sentence-transformers; "
        "считается при ingestion (батч), используется cosine-компонентом"
    )
    b.bullet("source_url, event_type, target_audience — справочные поля для UI")

    b.h3("Откуда берётся событие")
    b.p(
        "Шесть ingestion-адаптеров с единым интерфейсом: Habr (BeautifulSoup), "
        "RSS/Atom (feedparser), KudaGo (публичный JSON API), Lu.ma (ICS), "
        "Meetup (GraphQL, токен), Telegram-каналы (web-preview scraping). "
        "Сырьё → raw_events (status=raw). EventNormalizerAgent (одна нода "
        "LangGraph + один LLM-вызов Groq) возвращает структурированный JSON; "
        "Pydantic валидирует; не-IT события (topics == []) → status=non_it."
    )
    b.p(
        "Перед сохранением — семантическая дедупликация: ищем существующее "
        "событие с cosine ≥ 0.92 за последние 90 дней. Если найдено — "
        "новое не создаётся."
    )

    # 7.3 9 компонент
    b.h2("7.3. Девять компонент hybrid-скоринга — каждый отдельно")
    b.p(
        "Каждый компонент — отдельная функция _component_* в "
        "app/recommender/hybrid.py. Возвращает float (вклад) или None "
        "(если сигнал не сработал). Веса — score_*_weight в "
        "app/core/config.py. Сбой одного компонента не ломает hybrid."
    )

    def comp(idx: int, name: str, formula: str, inputs: str, weight: str, role: str):
        b.h3(f"Компонент {idx}. {name}")
        b.p(f"Формула / суть: {formula}")
        b.p(f"Входы: {inputs}")
        b.p(f"Вес в config.py: {weight}")
        b.p(f"Что добавляет в выдачу: {role}")

    comp(
        1, "rule (правила профиля)",
        "+2 за пересечение тем, +3 за preferred_format=event.format (+1 если any), "
        "+2 за совпадение города (+1 если any). Умножается на score_rule_weight.",
        "users.topic_weights, preferred_format, city; events.topics, format, city.",
        "score_rule_weight = 0.5",
        "Надёжный baseline без LLM. Гарантирует осмысленную выдачу, если "
        "остальные компоненты упадут.",
    )
    comp(
        2, "cosine (семантическая близость)",
        "cosine(user_emb, event_emb) × score_cosine_weight. user_emb — смесь "
        "70% профильного и 30% session-emb последних 5 взаимодействий.",
        "users.embedding (кэш), events.embedding (кэш), последние 5 interactions.",
        "score_cosine_weight = 10.0",
        "Близость по смыслу: «нейросети для текста» найдёт «NLP с LLM».",
    )
    comp(
        3, "bayesian (Thompson sampling)",
        "Для каждой темы события — p ~ Beta(α, β) с применённым decay, "
        "берётся max(p) по темам, умножается на score_bayes_weight.",
        "user_topic_stats, events.topics.",
        "score_bayes_weight = 5.0",
        "Долгосрочная адаптация под темы + автоматический explore/exploit.",
    )
    comp(
        4, "quality (качество события)",
        "(event.quality_score / 10.0) × score_quality_weight.",
        "events.quality_score (LLM при нормализации).",
        "score_quality_weight = 0.5",
        "Поднимает содержательное, опускает «вебинары-однодневки».",
    )
    comp(
        5, "hype (хайп / актуальность)",
        "(event.hype_score / 10.0) × score_hype_weight.",
        "events.hype_score (LLM при нормализации).",
        "score_hype_weight = 0.3",
        "Громкие конференции выигрывают у нишевых при прочих равных.",
    )
    comp(
        6, "freshness (свежесть даты)",
        "exp(−ln(2) × |days| / freshness_half_life_days) × score_freshness_weight. "
        "half_life_days = 30 по умолчанию.",
        "events.date, settings.freshness_half_life_days.",
        "score_freshness_weight = 2.0",
        "Прошедшие и далёкие-будущие события мягко уходят в хвост.",
    )
    comp(
        7, "skill_gap (соответствие карьерным целям)",
        "|target_skills ∩ event.tech_stack| / |target_skills| + бонус за seniority. "
        "× score_skill_gap_weight.",
        "user_skill_profiles, events.tech_stack, events.seniority.",
        "score_skill_gap_weight = 3.0",
        "Превращает рекомендатель из «интересного» в «приближающее к цели». "
        "Работает только после /bio cold-start.",
    )
    comp(
        8, "bandit (LinUCB)",
        "UCB: xᵀθ̂ + α·√(xᵀA⁻¹x), где θ̂ = A⁻¹·b. x — 16-мерный контекстный "
        "вектор пары (user, event). × bandit_weight.",
        "user_bandit_states (A, b).",
        "bandit_weight = 2.0",
        "Контекстная online-персонализация + exploration-бонус событиям с "
        "непривычным контекстом.",
    )
    comp(
        9, "gnn (LightGCN коллаборативный)",
        "user_emb · event_emb из натренированной LightGCN-модели. × gnn_weight.",
        "data/gnn_embeddings.npz (от scripts/train_gnn.py).",
        "gnn_weight = 3.0; GNN_ENABLED=false по умолчанию",
        "«Похожие на тебя лайкали вот это». На малом датасете выключен; "
        "включается явно после тренировки.",
    )

    # 7.4 Как вместе
    b.h2("7.4. Как 9 компонент работают вместе")

    b.h3("Шаг 1. Подсчёт hybrid_score")
    b.p(
        "На /recommend бэкенд достаёт все события и для каждого вызывает "
        "compute_score_breakdown(user, event, db). Внутри последовательно "
        "считаются 9 компонент, каждая в try/except. Сбой → 0.0 в breakdown "
        "+ WARNING в логи. Итоговый score = сумма breakdown."
    )

    b.h3("Шаг 2. Сортировка и MMR")
    b.p(
        "Сортировка по score убыванию. Поверх — MMR (λ=0.7): каждое "
        "следующее событие должно быть и релевантным, и не слишком похожим "
        "на уже выбранные. Без MMR в топ-5 иногда попадали 5 митапов про "
        "Kubernetes."
    )

    b.h3("Шаг 3. Передача в бот")
    b.p(
        "API возвращает JSON: title, description, score, score_breakdown, "
        "explanation. Бот рендерит первую карточку и кнопки ❓ Почему / "
        "📖 Подробнее / 👍 / 👎 / ⭐ / Похожие / Следующее. Полный "
        "explanation — по кнопке, а не в самой карточке."
    )

    b.h3("Шаг 4. Обратная связь")
    b.p("Каждый лайк/дизлайк/save обновляет четыре подсистемы в одной транзакции:")
    b.bullet("topic_weights в users (+3 / −2 / +1 по темам события)")
    b.bullet("alpha/beta в user_topic_stats (соответственно)")
    b.bullet("Матрица A и вектор b в user_bandit_states (LinUCB)")
    b.bullet(
        "Long-term memory — best-effort: LLM пишет заметку, если фидбэк "
        "значимый. Изолировано в SAVEPOINT."
    )
    b.p(
        "Команда /undo делает rollback во ВСЕХ четырёх подсистемах с "
        "direction=-1 — в большинстве рекомендательных систем бандит «не "
        "умеет забывать», у нас умеет."
    )

    b.h3("Шаг 5. Объяснимость")
    b.p(
        "Поскольку каждый компонент возвращает свой вклад, мы знаем «кто "
        "сколько весит». Из этого две вещи:"
    )
    b.bullet(
        "❓ Почему — короткий ответ LLM в поп-апе (1–2 предложения, без "
        "цифр, с привязкой к одному конкретному факту: тема, target_skill, "
        "лайкнутое событие)."
    )
    b.bullet(
        "📖 Подробнее — отдельное сообщение: LLM-нарратив + 1–2 совета + "
        "rule-based разбивка по компонентам с числами + counterfactual "
        "(«без skill_gap-сигнала событие потеряло бы ~45 % веса»)."
    )

    b.h3("Шаг 6. Контекстная и коллаборативная парадигмы — в одной системе")
    b.bullet(
        "КОНТЕКСТНАЯ — компонент bandit (LinUCB) считает рекомендации с "
        "учётом 16-мерного контекстного вектора пары. Per-user A и b "
        "обновляются онлайн на каждый фидбэк. Classic contextual bandit."
    )
    b.bullet(
        "КОЛЛАБОРАТИВНАЯ — компонент gnn (LightGCN) строит двудольный граф "
        "user↔event по лайкам/saves и через пропагацию по графу получает "
        "эмбеддинги, где похожие пользователи и события рядом. Чистый numpy, "
        "без torch-geometric. По умолчанию выключен — мал датасет."
    )
    b.p(
        "Помимо этих двух, в hybrid есть content-based (cosine), rule-based "
        "(rule), Bayesian preference (bayesian) и эвристики (quality, hype, "
        "freshness, skill_gap). Hybrid = композиция четырёх парадигм + три "
        "эвристики."
    )

    b.h3("Сводка: 9 компонент в одной таблице")
    table = b.doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Компонент"
    hdr[1].text = "Парадигма"
    hdr[2].text = "Источник данных"
    hdr[3].text = "Что добавляет"
    rows = [
        ("rule", "Rule-based", "Профиль и темы события", "Фолбэк без AI"),
        ("cosine", "Content-based", "Embeddings", "Близость по смыслу"),
        ("bayesian", "Bayesian preference", "α/β по темам", "Адаптация + explore/exploit"),
        ("quality", "Эвристика качества", "LLM-оценка", "Содержательные в топ"),
        ("hype", "Эвристика хайпа", "LLM-оценка", "Громкие события в топ"),
        ("freshness", "Эвристика времени", "Дата события", "Прошедшие/далёкие вниз"),
        ("skill_gap", "Goal-aware", "Skill-профиль + tech_stack", "Под карьерные цели"),
        ("bandit", "Contextual bandit", "Матрица A, вектор b", "Контекстная персонализация"),
        ("gnn", "Collaborative", "Граф user↔event", "Похожие на тебя лайкали"),
    ]
    for r in rows:
        cells = table.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v


# ════════════════════════════════════════════════════════════════════════
#  План показа курсовой работы
# ════════════════════════════════════════════════════════════════════════


def build_presentation_plan(b: Builder) -> None:
    b.p("План показа курсовой работы")
    b.p("EventMind — агрегация IT-событий и персонализированные AI-рекомендации")
    b.p(
        "Подробная версия с объяснением всех терминов · ориентир: 15–20 минут "
        "на защиту"
    )
    b.empty()
    b.p(
        "Документ написан в расчёте на читателя, который видит проект "
        "впервые и не является профильным специалистом по recsys. Для каждой "
        "технологии — пояснение «что это» и «зачем здесь»."
    )
    b.empty()

    # ── 1 ────────────────────────────────────────────────────────────────
    b.h1("1. Что такое EventMind в одну минуту")
    b.p(
        "EventMind — сервис-помощник для IT-специалистов. Автоматически "
        "собирает анонсы IT-мероприятий из 6 источников, нормализует через "
        "LLM, строит персональную ленту с объяснением «почему именно это» "
        "и даёт мульти-турового AI-копилота, который умеет планировать "
        "карьерный рост и подбирать события под цель."
    )
    b.h3("Какую реальную проблему решает")
    b.bullet(
        "Проблема: IT-мероприятий очень много, анонсы разбросаны по десятку "
        "сайтов и Telegram-каналов, отфильтровать релевантное по уровню/"
        "формату/городу — рутина."
    )
    b.bullet(
        "Решение: один Telegram-бот собирает события со всех источников, "
        "понимает уровень, темы и стек, и строит ленту под каждого "
        "пользователя по 9 признакам."
    )
    b.bullet(
        "Главное отличие от «угадай, что мне понравится»: мы знаем не только "
        "темы пользователя, но и его карьерную цель и историю реакций, и "
        "поднимаем события, которые приближают к цели."
    )
    b.h3("Что увидит преподаватель на демо")
    b.bullet("Telegram-бот: настройка профиля (3 шага), reply-меню")
    b.bullet("Лента рекомендаций с кнопками ❓ Почему / 📖 Подробнее / 👍/👎/⭐")
    b.bullet("Лайк/дизлайк → лента подстраивается в реальном времени")
    b.bullet("AI-копилот: формирует план развития на 6 месяцев под цель")
    b.bullet("Админ-эндпоинты — состояние БД и подсистем (опционально)")
    b.empty()

    # ── 2 Глоссарий ──────────────────────────────────────────────────────
    b.h1("2. Глоссарий: все термины простыми словами")
    b.p(
        "Главный раздел. Если читать только его — уже можно вести "
        "осмысленный разговор по проекту. Все следующие разделы ссылаются "
        "на термины отсюда."
    )

    b.h2("2.1. Базовые термины IT")

    def term(name: str, simple: str, why: str, soundbite: str):
        b.h3(name)
        b.p(f"Простыми словами: {simple}")
        b.p(f"Зачем это в проекте: {why}")
        b.p(f"Если спросят — можно сказать так: «{soundbite}»")

    term(
        "Python",
        "Язык программирования, на котором написан весь проект. Версия 3.12.",
        "Самый популярный язык для AI/ML и веб-бэкенда, много готовых библиотек.",
        "Весь backend и AI-логика — на Python 3.12.",
    )
    term(
        "Backend / Бэкенд",
        "Невидимая «серверная» часть программы. Принимает запросы, работает "
        "с базой, считает рекомендации, отдаёт результат боту.",
        "Без бэкенда нечем хранить пользователей и события, нечем считать "
        "рекомендации.",
        "Бэкенд — это сердце системы, он принимает запросы и отвечает "
        "JSON'ом.",
    )
    term(
        "API",
        "Соглашение, как одна программа общается с другой: какие URL "
        "вызывать, какие данные передавать, что вернётся.",
        "Telegram-бот разговаривает с бэкендом через API — это разделяет "
        "ответственность.",
        "Бот — клиент, API — сервер. Бот ничего не знает про БД, только "
        "про URL.",
    )
    term(
        "REST",
        "Стиль API на HTTP: GET для чтения, POST для записи, чистые URL.",
        "FastAPI ровно про REST: все эндпоинты — обычные HTTP, легко "
        "тестировать curl'ом.",
        "У нас классический REST: /recommendations/{id}, /events/search, "
        "/copilot/{id}/turn.",
    )
    term(
        "JSON",
        "Формат обмена данными: фигурные скобки и пары ключ–значение.",
        "Все ответы API — JSON, бот парсит и рендерит.",
        "Бот получает JSON со списком карточек и рисует первую.",
    )
    term(
        "FastAPI",
        "Современный Python-фреймворк для REST API: декларативные роуты, "
        "автогенерация OpenAPI, нативная асинхронность.",
        "Удобнее Flask и быстрее Django — подходит для тонкого API-слоя.",
        "FastAPI отдаёт OpenAPI на /docs, видно все эндпоинты сразу.",
    )
    term(
        "uvicorn",
        "ASGI-сервер, который запускает FastAPI-приложение.",
        "uvicorn слушает порт и передаёт запросы FastAPI.",
        "Запускается командой uvicorn app.api.main:app.",
    )
    term(
        "Pydantic",
        "Библиотека для валидации данных: описываешь схему, и она проверяет "
        "входной JSON на типы и поля.",
        "Все входы/выходы API и AI-нормализатор проходят через Pydantic-"
        "схемы — это убирает классы багов «не то поле, не тот тип».",
        "LLM возвращает текст, Pydantic превращает его в типизированный "
        "объект — если LLM «солгал», ловим ValidationError.",
    )
    term(
        "База данных (БД)",
        "Программа, которая хранит данные на диске и умеет быстро искать.",
        "Храним пользователей, события, лайки, embedding'и, состояние "
        "бандита, заметки памяти.",
        "В dev — SQLite, в проде — PostgreSQL + pgvector для быстрого "
        "поиска по векторам.",
    )
    term(
        "ORM (SQLAlchemy)",
        "Слой между Python-кодом и SQL: классы вместо CREATE TABLE.",
        "Один и тот же код работает и с SQLite, и с Postgres — это важно "
        "для деплоя.",
        "У нас 11 моделей SQLAlchemy в app/db/models.",
    )
    term(
        "Миграции (Alembic)",
        "Версионирование схемы БД: каждое изменение — отдельный файл.",
        "10 миграций; включая pgvector, copilot_sessions, bandit_state, "
        "user_memories.",
        "alembic upgrade head накатывает все миграции по порядку.",
    )
    term(
        "Docker / docker-compose",
        "Способ упаковать сервис со всеми зависимостями в контейнер; "
        "compose описывает несколько контейнеров вместе.",
        "У нас три сервиса в compose: api, bot, scheduler. Healthcheck на "
        "api перед стартом bot и scheduler.",
        "docker-compose up — и всё поднимается одной командой.",
    )

    b.h2("2.2. Telegram и боты")
    term(
        "Telegram Bot API",
        "HTTPS-интерфейс Telegram для ботов: отправлять сообщения, "
        "получать апдейты, рисовать кнопки.",
        "Без него никаких ботов в TG не существует.",
        "Все сообщения наш бот шлёт через POST на api.telegram.org/botXXX.",
    )
    term(
        "aiogram 3",
        "Python-фреймворк поверх Bot API: декораторы для команд, FSM, "
        "роутеры, асинхронность.",
        "Удобнее писать обработчики команд, чем напрямую POST'ить.",
        "У нас 6 router'ов: start, recommendations, profile, search, "
        "copilot, subscriptions.",
    )
    term(
        "Long-polling",
        "Бот периодически спрашивает Telegram: «есть новые сообщения?» — "
        "альтернатива вебхукам.",
        "Не нужен публичный URL — удобнее для разработки и деплоя на любой "
        "VPS.",
        "У нас dp.start_polling(bot) — без вебхуков.",
    )
    term(
        "Inline-кнопки",
        "Кнопки, прикреплённые к сообщению — при нажатии бот получает "
        "callback_data.",
        "Все действия в карточке события — inline-кнопки: ❓ Почему, "
        "👍/👎/⭐, Похожие, Следующее.",
        "callback_data — это строка, мы парсим её регулярками в роутере.",
    )
    term(
        "Reply-кнопки",
        "Кнопки в нижнем меню, как будто пользователь сам пишет текст.",
        "У нас 4 reply-кнопки главного меню: «🎯 Рекомендации», "
        "«🔍 Поиск», «👤 Профиль», «⚙️ Ещё».",
        "Reply-кнопка отправляет в чат текст «🎯 Рекомендации», а handler "
        "ловит F.text == «🎯 Рекомендации».",
    )

    b.h2("2.3. Искусственный интеллект (LLM, эмбеддинги, RAG)")
    term(
        "LLM (Large Language Model)",
        "Большая нейросеть, которая «продолжает» текст: ChatGPT — самый "
        "известный пример.",
        "Используется в трёх местах: нормализация событий, объяснения "
        "«Почему», мульти-агентный Copilot.",
        "Цепочка _LLMChain: Gemini (через REST-транспорт, автопроба модели) → "
        "Groq llama-3.3-70b → Groq llama-3.1-8b — последовательный fallback "
        "на rate-limit/сбое.",
    )
    term(
        "Groq",
        "Облачный провайдер LLM с очень быстрым инференсом (специальное "
        "железо LPU).",
        "Дешевле и быстрее OpenAI на задачах нашего масштаба.",
        "Один вызов нормализации события — ~0.5 с против ~3 с у конкурентов.",
    )
    term(
        "Embedding (векторное представление)",
        "Кодируем текст в 384-мерный вектор так, что близкие по смыслу "
        "тексты дают близкие векторы.",
        "Сравниваем профиль пользователя и описание события через cosine — "
        "это даёт «понимание смысла», а не только точные слова.",
        "У нас sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2, "
        "поддерживает русский.",
    )
    term(
        "Cosine similarity",
        "Метрика близости двух векторов: 1.0 — одинаковые, 0 — независимые, "
        "−1 — противоположные.",
        "Главная метрика семантической близости в проекте: между "
        "user_embedding и event_embedding, между событиями (для MMR), "
        "между memory-заметками.",
        "Если cosine ≥ 0.92 за 90 дней — мы считаем события дубликатами.",
    )
    term(
        "pgvector",
        "Расширение PostgreSQL: тип данных vector + индексы + оператор <=> "
        "для cosine-distance.",
        "На проде позволяет искать top-k по семантике на стороне БД, а не "
        "сортировать в Python.",
        "Включается миграцией 9a1b2c3d4e5f; на SQLite остаётся JSON в TEXT.",
    )
    term(
        "RAG (Retrieval-Augmented Generation)",
        "Сначала найти релевантные документы (retrieve), потом скормить их "
        "LLM в промпт (generate).",
        "Все наши AI-объяснения — RAG: hybrid отбирает top-N, LLM пишет "
        "пояснения только по ним. Без RAG модель бы фантазировала.",
        "Запрос /recommend → hybrid отбирает 5 событий → LLM объясняет "
        "только эти 5 — она не «придумывает» событий.",
    )
    term(
        "Function calling (вызов инструментов)",
        "LLM умеет сама решать «нужно вызвать функцию foo с такими "
        "аргументами», мы возвращаем результат, она продолжает.",
        "У Copilot 6 tools: search_events, get_user_profile, "
        "get_interactions_summary, explain_event, recall_about_user, mark_saved.",
        "Если пользователь спрашивает «что я лайкал на прошлой неделе», "
        "Copilot сам вызовет get_interactions_summary.",
    )
    term(
        "LangGraph",
        "Фреймворк для построения multi-agent workflow в виде графа: "
        "ноды + рёбра + условные переходы.",
        "Copilot собран как граф: retrieve → supervisor → один из 5 "
        "специалистов → finalize.",
        "LangGraph лучше LangChain Agents: явный граф вместо «магической» "
        "цепочки.",
    )
    term(
        "Supervisor-Worker",
        "Архитектурный паттерн: один агент-координатор решает, какому "
        "специалисту передать задачу.",
        "Supervisor классифицирует intent (recommend/career/roadmap/"
        "explain/summary) и роутит на нужную ноду.",
        "Это позволяет иметь короткие специализированные промпты вместо "
        "одного гигантского.",
    )
    term(
        "Hallucination (галлюцинации LLM)",
        "LLM «придумывает» факты, которых не было в данных.",
        "У нас в системных промптах объяснений запрет на выдуманных "
        "спикеров, программу, дату — модель обязана сослаться на конкретный "
        "факт из profile/event.",
        "Pydantic-схема выходов + явные «не фантазируй» в промптах + RAG "
        "= минимум галлюцинаций.",
    )

    b.h2("2.4. Рекомендательная система — термины")
    term(
        "Hybrid recommender",
        "Рекомендатель, который комбинирует несколько разных алгоритмов "
        "(rule + content + collaborative + …) в один score.",
        "У нас 9 компонент, итог — взвешенная сумма. Преимущества "
        "каждого подхода складываются, недостатки взаимно компенсируются.",
        "Если LLM упадёт — rule всё ещё работает. Если истории мало — "
        "content-based закрывает. Если граф плотный — gnn даёт серьёзный "
        "буст.",
    )
    term(
        "Rule-based scoring",
        "Простые правила: «совпала тема — +2 балла, совпал формат — +3».",
        "Самый надёжный компонент: работает без AI, без БД-кэша, без "
        "истории. Это наш fallback по умолчанию.",
        "Если все остальные компоненты упадут — лента всё равно остаётся "
        "осмысленной за счёт rule.",
    )
    term(
        "Topic weights (веса тем)",
        "JSON-словарь {тема: вес} в profiles. По умолчанию вес=5, лайк +3, "
        "save +1, дизлайк −2.",
        "Личный «отпечаток интересов» на длительный период.",
        "Веса используются в rule-компоненте: пересечения с topic_weights "
        "> 3 дают бонус.",
    )
    term(
        "Bayesian Thompson sampling",
        "На каждую тему — параметры Beta-распределения (α, β). При "
        "ранжировании сэмплируем p ~ Beta(α, β) — это даёт автоматический "
        "баланс «использовать известное» и «пробовать новое».",
        "Альтернатива жёстким коэффициентам exploration — самоадаптивная "
        "и более robust.",
        "После 10–20 фидбэков лента заметно «подстраивается» под "
        "конкретного пользователя.",
    )
    term(
        "Contextual bandit (LinUCB)",
        "Онлайн-обучаемая линейная регрессия с UCB-бонусом за "
        "неопределённость. Обновляется на каждый фидбэк.",
        "Учитывает не отдельную тему, а комбинацию признаков "
        "(тема + формат + город + freshness).",
        "Это classic contextual bandit из RL. У нас на чистом numpy, без "
        "torch.",
    )
    term(
        "LightGCN (collaborative)",
        "Графовая нейросеть: двудольный граф user↔event по лайкам, "
        "пропагация по графу даёт эмбеддинги.",
        "Похожие пользователи + похожие события оказываются рядом в "
        "пространстве — это коллаборативный сигнал.",
        "Реализован, но выключен на маленьком датасете прототипа: граф "
        "слишком разрежен и сигнал шумит.",
    )
    term(
        "MMR (Maximal Marginal Relevance)",
        "Пост-процессинг: каждое следующее событие должно быть и "
        "релевантным, и не слишком похожим на уже выбранные.",
        "Без MMR в топ-5 иногда попадали 5 митапов про Kubernetes — теперь "
        "разнообразие гарантируется.",
        "λ=0.7 — компромисс. 1.0 = чистая релевантность, 0.0 = чистое "
        "разнообразие.",
    )
    term(
        "Freshness (свежесть)",
        "Экспоненциальный decay по дате события: 1.0 на сегодня, плавно "
        "уменьшается в обе стороны.",
        "Прошедшие и далёкие-будущие события не доминируют в выдаче.",
        "Half-life = 30 дней. За месяц вес падает вдвое.",
    )
    term(
        "Quality / Hype",
        "Два независимых скора 1–10 от LLM при нормализации.",
        "Quality — содержательность, hype — громкость/актуальность. Учат "
        "разному.",
        "Без quality в топ иногда попадали сомнительные «вебинары-"
        "однодневки»; добавили — стало чище.",
    )
    term(
        "Skill-gap",
        "Пересечение целевых навыков пользователя с tech_stack события + "
        "бонус за совпадение seniority.",
        "Поднимает события, которые приближают пользователя к карьерной "
        "цели.",
        "Работает только если пользователь прошёл /bio cold-start (LLM "
        "извлекает target_skills).",
    )
    term(
        "Cold start (холодный старт)",
        "Проблема: новый пользователь без истории — рекомендатель не знает, "
        "что ему нравится.",
        "Решение: /bio — пользователь пишет «я backend, опыт 3 года, хочу "
        "в ML», LLM с pydantic извлекает skill-профиль и виртуально "
        "обновляет topic_weights.",
        "После /bio даже без единого лайка лента уже подстроена.",
    )
    term(
        "Long-term memory (mem0-style)",
        "Активные заметки про пользователя от первого лица: «хочет в ML», "
        "«не любит offline». Пишет LLM-агент сам.",
        "Используется Copilot'ом через recall_about_user — это даёт "
        "ощущение «AI меня помнит».",
        "При накоплении >50 заметок — LLM группирует и сжимает до ≤5 на "
        "категорию.",
    )

    b.h2("2.5. Источники событий")
    term(
        "Ingestion",
        "Процесс загрузки и обработки внешних данных.",
        "У нас 6 ingestion-адаптеров с единым интерфейсом.",
        "Habr, RSS, KudaGo, Lu.ma, Meetup, Telegram-каналы — каждый в "
        "отдельном модуле app/ingestion/sources.",
    )
    term(
        "BeautifulSoup",
        "Библиотека для парсинга HTML.",
        "Habr-адаптер тянет страницу анонса и достаёт title, body, дату.",
        "Используется только в habr.py — для остальных есть feeds/API.",
    )
    term(
        "RSS / Atom",
        "Старый-добрый формат ленты новостей: XML с новыми записями.",
        "Унифицированный feedparser-адаптер работает с любым RSS-feed'ом.",
        "В .env RSS_FEEDS=url1,url2,url3 — добавлять источники без кода.",
    )
    term(
        "Telegram-каналы (web-preview)",
        "У публичного канала есть t.me/s/<имя> — там в HTML видны "
        "последние сообщения.",
        "Не требует api_id/api_hash, только публичный username канала.",
        "TG_INGEST_CHANNELS=iteventsru,ITMeeting — два проверенных канала.",
    )
    term(
        "Нормализация (raw_event → event)",
        "AI-агент извлекает из сырого текста структурированные поля: "
        "title, format, city, level, topics, tech_stack, quality, hype.",
        "Без этого мы бы хранили «ёлки-палки» из разных источников.",
        "Один LLM-вызов на событие, валидация Pydantic, не-IT отсеиваются "
        "с status=non_it.",
    )
    term(
        "Семантическая дедупликация",
        "Если новое событие имеет cosine ≥ 0.92 с уже сохранённым за "
        "последние 90 дней — не сохраняем.",
        "Ловит парафразы: один и тот же митап от Habr и Telegram-канала.",
        "Порог настраивается через DEDUP_THRESHOLD; отключается "
        "DEDUP_ENABLED=false.",
    )

    b.h2("2.6. Прочие термины проекта")
    term(
        "APScheduler",
        "Планировщик задач: «каждые 6 часов делать X».",
        "В отдельном процессе крутятся 4 джоба: digest, ingest_habr, "
        "ingest_rss, ingest_telegram + compact_memories раз в 7 дней.",
        "Интервалы конфигурируются через .env: INGEST_INTERVAL_HOURS, "
        "DIGEST_INTERVAL_MINUTES.",
    )
    term(
        "Healthcheck",
        "Простой эндпоинт, который говорит «жив или нет».",
        "В docker-compose bot и scheduler ждут /health api перед стартом.",
        "/health возвращает {db, llm, ingestion} — состояние каждой "
        "подсистемы.",
    )
    term(
        "Fallback",
        "Запасной путь при сбое основного.",
        "У нас три уровня: LLM → groq_fallback_model, /recommendations → "
        "rule-only при сбое hybrid, Copilot → «AI временно недоступен» с "
        "сохранением сессии.",
        "Цель: пользователь никогда не видит «ничего не работает» — всегда "
        "хоть какая-то осмысленная выдача.",
    )
    term(
        "Pytest",
        "Фреймворк для тестов в Python: функции с assert + фикстуры.",
        "180 тестов в 25 файлах: scoring, bayesian, bandit, gnn, memory, "
        "ingestion, supervisor, copilot tools, cold-start, scheduler.",
        "pytest -q проходит за ~30 с.",
    )
    term(
        "Ruff",
        "Линтер для Python: быстрая проверка стиля и типичных ошибок.",
        "0 ошибок на app/ и tests/.",
        "ruff check . — за секунду.",
    )
    b.empty()

    # ── 3 ────────────────────────────────────────────────────────────────
    b.h1("3. Общая архитектура: как всё устроено")
    b.h2("3.1. Три процесса")
    b.bullet("api (FastAPI + uvicorn) — REST-сервер, точка входа всех запросов")
    b.bullet("bot (aiogram 3) — Telegram-клиент, long-polling, общается с api по HTTP")
    b.bullet(
        "scheduler (APScheduler) — фоновый процесс: AI-дайджест + периодический "
        "ingestion + compaction памяти"
    )
    b.p("Все три в одном docker-compose, БД — общая (SQLite файл / Postgres сервис).")

    b.h2("3.2. Поток данных")
    b.bullet("scheduler → ingestion-источники → raw_events → AI-нормализатор → events")
    b.bullet("user в Telegram → bot → API /recommendations → hybrid → MMR → bot → рисует карточку")
    b.bullet("user жмёт 👍 → bot → API /interactions → 4 подсистемы (weights, bayesian, bandit, memory) обновляются")
    b.bullet("user пишет /copilot цель → API /copilot/turn → LangGraph (retrieve → supervisor → специалист → finalize) → bot")

    b.h2("3.3. Готовые диаграммы")
    b.bullet("docs/diagrams/architecture.png — три процесса и потоки")
    b.bullet("docs/diagrams/recsys_pipeline.png — 9 компонент → сортировка → MMR")
    b.bullet("docs/diagrams/copilot_graph.png — Supervisor-Worker граф")
    b.bullet("docs/diagrams/er.png — таблицы и связи")
    b.bullet("docs/diagrams/ingestion.png — 6 источников → нормализация → events")
    b.empty()

    # ── 4 ────────────────────────────────────────────────────────────────
    b.h1("4. Технологический стек — что и зачем")
    b.bullet("Python 3.12, FastAPI, aiogram 3, SQLAlchemy 2 + Alembic, Pydantic 2")
    b.bullet("LangGraph 0.2 + multi-provider _LLMChain: Gemini (REST + autopobe) → Groq 70b → Groq 8b")
    b.bullet("sentence-transformers 2.7 (MiniLM-L12-v2, 384 dim, поддерживает русский)")
    b.bullet("APScheduler 3.10 (background jobs)")
    b.bullet("feedparser, BeautifulSoup 4 + lxml, httpx")
    b.bullet("PostgreSQL + pgvector (prod) / SQLite + WAL (dev)")
    b.bullet("Docker + docker-compose; pytest 8.2, ruff")
    b.empty()

    # ── 5 ────────────────────────────────────────────────────────────────
    b.h1("5. Структура базы данных — что и зачем хранится")
    b.bullet("users — профиль (telegram_id, format, city, topic_weights JSON, embedding)")
    b.bullet("topics — справочник тем; events — нормализованные события (title, format, city, level, date, quality, hype, tech_stack JSON, embedding)")
    b.bullet("user_topics, event_topics — many-to-many связи")
    b.bullet("interactions — лайки/дизлайки/saves")
    b.bullet("user_topic_stats — Bayesian α/β по темам")
    b.bullet("user_bandit_states — матрица A и вектор b LinUCB")
    b.bullet("user_skill_profiles — current/target роль, скиллы, цели")
    b.bullet("user_memories — заметки long-term памяти")
    b.bullet("copilot_sessions — история мульти-туровых диалогов")
    b.bullet("raw_events — сырьё ingestion со статусом raw/normalized/non_it")
    b.h3("Как объяснить структуру преподавателю")
    b.p(
        "Три кластера таблиц: «события» (events + raw_events + topics + "
        "event_topics), «пользователь» (users + 5 sidecar-таблиц для модели "
        "пользователя), «взаимодействия» (interactions + copilot_sessions). "
        "Каждая sidecar-таблица — отдельная подсистема: bayesian, bandit, "
        "skill, memory."
    )
    b.empty()

    # ── 6 ────────────────────────────────────────────────────────────────
    b.h1("6. Как работает рекомендация — пошагово")
    b.h2("6.1. Что происходит на /recommend")
    b.bullet("Бэкенд достаёт все события (db.query(Event).all())")
    b.bullet("Прогрев: refresh_user_embedding + ensure_event_embeddings (батчем недостающие)")
    b.bullet("Для каждого события считаются 9 компонент в compute_score_breakdown")
    b.bullet("Каждая в try/except: сбой → 0.0 + WARNING")
    b.bullet("Сортировка по сумме breakdown по убыванию")
    b.bullet("MMR-rerank: каждое следующее должно отличаться по эмбеддингу")
    b.bullet("Возврат JSON: title, description, score, score_breakdown, explanation")
    b.bullet(
        "ВАЖНО: при сбое flush embedding-кэша теперь делается rollback, "
        "сессия остаётся живой, ответ не падает 500-кой (это починено 29 мая)"
    )

    b.h2("6.2. Девять компонент hybrid-скоринга")
    b.p("Подробно — в Обзоре, глава 7. Кратко:")
    b.bullet("rule — правила профиля (надёжный baseline)")
    b.bullet("cosine — семантическая близость user_emb ↔ event_emb")
    b.bullet("bayesian — Thompson sampling по темам с temporal decay")
    b.bullet("quality — оценка содержательности (LLM)")
    b.bullet("hype — оценка хайповости (LLM)")
    b.bullet("freshness — экспоненциальный decay по дате")
    b.bullet("skill_gap — пересечение target_skills и tech_stack")
    b.bullet("bandit — LinUCB UCB-формула, онлайн-обучение")
    b.bullet("gnn — LightGCN скалярное произведение (выключен по умолчанию)")

    b.h2("6.3. Что происходит при 👍 (лайке)")
    b.bullet("POST /recommendations/interactions с action=like")
    b.bullet("В одной транзакции обновляются:")
    b.bullet("  users.topic_weights — +3 по каждой теме события")
    b.bullet("  user_topic_stats — alpha += 3 по каждой теме")
    b.bullet("  user_bandit_states — A += x·xᵀ, b += 1·x")
    b.bullet("  user_memories — LLM-агент может записать заметку (best-effort, в SAVEPOINT)")

    b.h2("6.4. Команда /undo — почему это нетривиально")
    b.p(
        "В большинстве рекомендеров «откатить лайк» = удалить запись. У нас "
        "лайк трогает 4 подсистемы — нужно везде сделать direction=-1: "
        "веса вернуть, α/β вернуть, A и b пересчитать с обратной x. "
        "create_interaction уже умеет toggle off — undo сводится к одному "
        "вызову с последним interaction'ом."
    )
    b.empty()

    # ── 7 ────────────────────────────────────────────────────────────────
    b.h1("7. AI-копилот — главная фича")
    b.h2("7.1. Что это такое в одном абзаце")
    b.p(
        "Мульти-туровый AI-помощник в Telegram. Пользователь пишет цель "
        "(«хочу за полгода стать middle backend»), бот за один проход "
        "анализирует профиль + историю + долговременную память + текущую "
        "ленту событий и возвращает план развития + 3–5 событий, которые "
        "его приближают. Дальше пользователь может уточнять «а если без "
        "конференций?» — Copilot помнит контекст. Сессия живёт 15 минут "
        "после последнего сообщения."
    )

    b.h2("7.2. Архитектура: Supervisor-Worker")
    b.p(
        "Граф LangGraph с пятью нодами специалистов и нодой-супервизором. "
        "Pipeline: retrieve → supervisor → conditional_edge → один из "
        "5 специалистов → finalize → END. Supervisor классифицирует "
        "intent через structured LLM-вызов (pydantic-валидация + keyword "
        "fallback)."
    )

    b.h2("7.3. Пять специалистов")
    b.bullet("recommendation_specialist — intent=recommend, подбор событий")
    b.bullet("career_coach — intent=career, анализ карьеры")
    b.bullet("roadmap_planner — intent=roadmap, план на 3/6/12 месяцев")
    b.bullet("event_explainer — intent=explain, разбор конкретного события")
    b.bullet("summary_specialist — intent=summary, сводки/обзоры")

    b.h2("7.4. Шесть инструментов (function calling)")
    b.bullet("search_events — семантический поиск событий по запросу")
    b.bullet("get_user_profile — снимок профиля + топ-темы")
    b.bullet("get_interactions_summary — сводка истории (лайки/сейвы/топ-темы)")
    b.bullet("explain_event — детальное объяснение, почему событие подходит")
    b.bullet("recall_about_user — top-k заметки long-term памяти")
    b.bullet("mark_saved — пометить событие сохранённым")

    b.h2("7.5. Мульти-туровый диалог")
    b.p(
        "После первого ответа Copilot бот ставит per-user state с "
        "session_id и timestamp. Любое следующее текстовое сообщение в "
        "течение 15 минут — не команда, не reply-кнопка, не «Найти:...» — "
        "продолжает сессию. История сохраняется в copilot_sessions "
        "(messages_json + last_message_at). При новом /copilot — открывается "
        "новая сессия, старая закрывается."
    )

    b.h2("7.6. Long-term memory в работе")
    b.p(
        "Когда пользователь активно лайкает события про MLOps, агент-"
        "memory пишет заметку «активно интересуется MLOps» с salience 7. "
        "При следующем запросе «как мне расти как ML-инженер» tool "
        "recall_about_user поднимает эту заметку, и Copilot ссылается на "
        "неё: «учитывая твой устойчивый интерес к MLOps, я бы…». Это и "
        "есть ощущение «AI меня помнит»."
    )
    b.empty()

    # ── 8 ────────────────────────────────────────────────────────────────
    b.h1("8. Источники событий и нормализация (ingestion)")
    b.h2("8.1. Шесть источников")
    b.bullet("Habr — BeautifulSoup-парсинг анонсов")
    b.bullet("RSS / Atom — feedparser (любой URL из .env)")
    b.bullet("KudaGo — публичный JSON API")
    b.bullet("Lu.ma — ICS-фиды конкретных календарей")
    b.bullet("Meetup — GraphQL (требует OAuth Pro-токен)")
    b.bullet("Telegram-каналы — web-preview scraping (t.me/s/<имя>)")

    b.h2("8.2. Пайплайн нормализации")
    b.bullet("Источник → raw_events (status=raw, payload — сырой JSON/HTML)")
    b.bullet("EventNormalizerAgent: один LLM-вызов на raw_event")
    b.bullet("Pydantic-схема валидирует возвращённый JSON")
    b.bullet("Если topics == [] → status=non_it, в events не пишем")
    b.bullet("Семантическая дедупликация: cosine ≥ 0.92 за 90 дней → дубль")
    b.bullet("Сохраняем event с embedding'ом, status=normalized")

    b.h2("8.3. Автоматическое пополнение")
    b.p(
        "В scheduler-процессе три ingestion-джоба: ingest_habr, ingest_rss, "
        "ingest_telegram. Интервал — INGEST_INTERVAL_HOURS=6 по умолчанию. "
        "На старте все три запускаются сразу (next_run_time=now), потом "
        "повторяются. RSS и Telegram сдвинуты на 30/60 с, чтобы не "
        "стартовать одновременно с Habr."
    )
    b.empty()

    # ── 9 ────────────────────────────────────────────────────────────────
    b.h1("9. Подготовка к защите (за 30 минут до показа)")
    b.h3("Чек-лист окружения")
    b.bullet("docker-compose up -d  → api, bot, scheduler поднимаются")
    b.bullet("curl localhost:8000/health → {db: ok, llm: ok, ingestion: ok}")
    b.bullet("В Telegram /start → тур → настройка профиля → /recommend → видим ленту")
    b.bullet("Нажать 👍 на 2–3 событиях → /recommend ещё раз → лента изменилась")
    b.bullet("/copilot хочу за полгода стать middle backend → видим план + 3–5 событий")
    b.bullet("Прислать в чат ещё одно сообщение «а если без конференций?» → Copilot помнит")
    b.bullet("В docker logs scheduler — видно [ingest:habr], [ingest:rss], [scheduler]")
    b.bullet("Сделать DIGEST_RUN_ON_STARTUP=true, DIGEST_INTERVAL_MINUTES=10 → дайджест приходит")

    b.h3("Резервный план (если что-то отвалится)")
    b.bullet("Не работает Telegram? Демонстрируем по curl: /recommendations/{id}")
    b.bullet("Не работает Groq? Объясняем: hybrid даёт rule-only выдачу, объяснение упрощённое")
    b.bullet("Не работает БД? Из бэкапа eventmind.db.bak (170+ событий)")
    b.bullet("Не работает GNN? Это OK: gnn_enabled=False по умолчанию")
    b.empty()

    # ── 10 ───────────────────────────────────────────────────────────────
    b.h1("10. Структура выступления (тайминг)")
    b.bullet("1 мин — введение (что и зачем)")
    b.bullet("2 мин — архитектура и стек")
    b.bullet("2 мин — ingestion-демо")
    b.bullet("3 мин — Telegram-бот: профиль + рекомендации")
    b.bullet("2 мин — AI-рекомендации (hybrid + LLM-объяснения)")
    b.bullet("3 мин — AI Copilot Supervisor-Worker (главная фича)")
    b.bullet("1 мин — admin/analytics (опционально)")
    b.bullet("1 мин — качество (тесты, миграции, надёжность)")
    b.bullet("1 мин — итоги и честно «что улучшить»")
    b.empty()

    # ── 11 ───────────────────────────────────────────────────────────────
    b.h1("11. Подробные сценарии показа")
    b.h2("11.1. Введение (1 минута)")
    b.p(
        "«EventMind — сервис, который помогает IT-специалисту находить "
        "релевантные мероприятия, не тратя час на ручной поиск по 10 "
        "разрозненным каналам. Под капотом — гибридный рекомендатель из "
        "9 компонент, мульти-агентный AI-копилот для карьерных задач, "
        "сбор данных из 6 источников с AI-нормализацией. Всё локально "
        "на одной машине, в трёх Docker-контейнерах.»"
    )

    b.h2("11.2. Архитектура и стек (2 минуты)")
    b.bullet("Показать docs/diagrams/architecture.png")
    b.bullet("Три процесса: api / bot / scheduler — независимы, общаются через REST")
    b.bullet("FastAPI 0.110 + aiogram 3 + LangGraph + Postgres/SQLite — современный AI-стек")
    b.bullet("Сказать про fallback на каждом уровне (LLM → 8b-модель, hybrid → rule, Copilot → сохранение сессии)")

    b.h2("11.3. Демо: пайплайн ingestion (2 минуты)")
    b.bullet("Показать docs/diagrams/ingestion.png — 6 источников")
    b.bullet("curl POST /ingestion/load-habr?limit=5 → видно нормализованные события")
    b.bullet("Подсветить: AI-нормализатор делает quality/hype/topics за один LLM-вызов")
    b.bullet("Сказать про семантическую дедупликацию (cosine ≥ 0.92)")

    b.h2("11.4. Демо: Telegram-бот, профиль и рекомендации (3 минуты)")
    b.h3("11.4.1. Onboarding")
    b.bullet("/start → 4-экранный тур")
    b.bullet("Кнопка 🚀 «Начать настройку» → выбор тем → формат → город")
    b.bullet("«Профиль настроен» → видим главное меню")
    b.h3("11.4.2. Hybrid-рекомендации")
    b.bullet("«🎯 Рекомендации» → карточка события")
    b.bullet("Подсветить кнопку ❓ Почему → короткий ответ LLM в поп-апе")
    b.bullet("📖 Подробнее → развёрнутый разбор + counterfactual + советы")
    b.h3("11.4.3. Обратная связь")
    b.bullet("👍 на 2–3 событиях → следующее /recommend показывает другие приоритеты")
    b.bullet("/undo → последняя реакция откатилась во всех 4 подсистемах")
    b.h3("11.4.4. Семантический поиск и cold-start")
    b.bullet("«🔍 Поиск» → «хочу научиться деплоить микросервисы» → keyword + semantic в один проход")
    b.bullet("/bio «backend 3 года, хочу в ML» → cold-start заполняет skill-профиль")

    b.h2("11.5. AI-рекомендации: hybrid + LLM-объяснения (2 минуты)")
    b.bullet("В коде: app/api/services/ai_recommendation_service.py")
    b.bullet("hybrid отбирает top-N → один LLM-вызов с pydantic-схемой → объяснения батчем")
    b.bullet("Подсветить: без LangGraph для этой задачи — это всего ОДИН вызов")
    b.bullet("Параметр limit реально влияет на длину выдачи (раньше было жёстко 3)")

    b.h2("11.6. AI-Copilot: Supervisor-Worker (3 минуты)")
    b.bullet("Показать docs/diagrams/copilot_graph.png")
    b.h3("11.6.1. Roadmap")
    b.bullet("/copilot хочу за полгода вырасти в middle backend")
    b.bullet("Supervisor → roadmap_planner")
    b.bullet("План: что освоить за месяц 1, 2, …, 6 + список событий")
    b.h3("11.6.2. Career coach")
    b.bullet("Продолжить: «а если у меня уже есть Kubernetes?»")
    b.bullet("Copilot помнит контекст (multi-turn), уточняет план")
    b.h3("11.6.3. Объяснение события")
    b.bullet("«объясни мне, почему ты порекомендовал митап X»")
    b.bullet("Supervisor → event_explainer → tool explain_event → LLM объясняет")
    b.h3("Что подсветить про архитектуру Copilot'а")
    b.bullet("Граф LangGraph: явный, тестируемый, не «магия»")
    b.bullet("5 коротких специализированных промптов вместо одного огромного")
    b.bullet("5 function-calling tools — модель сама решает, что вызвать")
    b.bullet("Long-term memory: recall_about_user поднимает заметки из user_memories")

    b.h2("11.7. Бэк-офис: admin-эндпоинты (1 минута)")
    b.bullet("/admin/stats — счётчики (users, events, interactions, memories)")
    b.bullet("/analytics/trending — тренды за 7 дней")
    b.bullet("Опционально, если останется время")

    b.h2("11.8. Качество (1 минута)")
    b.bullet("180 тестов в 25 файлах: scoring, bayesian, bandit, gnn, memory, ingestion, supervisor")
    b.bullet("ruff 0 ошибок")
    b.bullet("10 миграций Alembic — schema versioning")
    b.bullet("3 уровня fallback: LLM, hybrid, Copilot")

    b.h2("11.9. Итоги (1 минута)")
    b.h3("Что сделано")
    b.bullet("Hybrid recommender из 9 компонент: rule + content + bayesian + контекстный bandit + collaborative gnn + эвристики")
    b.bullet("Multi-agent Copilot на LangGraph с 5 специалистами и 5 tools")
    b.bullet("6 ingestion-источников с AI-нормализацией")
    b.bullet("Long-term memory mem0-style")
    b.bullet("Объяснимость двух уровней (короткий поп-ап + развёрнутый разбор)")
    b.h3("Какие технологии освоены")
    b.bullet("LangGraph, function calling, structured output (pydantic)")
    b.bullet("Contextual bandit (LinUCB) с нуля на numpy")
    b.bullet("LightGCN с нуля на numpy")
    b.bullet("Bayesian Thompson sampling с temporal decay")
    b.bullet("pgvector + Postgres-индексы")
    b.h3("Что можно улучшить (честно)")
    b.bullet("GNN-граф пока шумит — нужен больший датасет")
    b.bullet("MMR-параметр λ подобран эмпирически — можно оптимизировать на данных")
    b.bullet("LLM-судья в эвале — пока однократно; нужен периодический прогон")
    b.bullet("На больших объёмах ingestion на SQLite иногда «database is locked» — фикс через PG")
    b.empty()

    # ── 12 ───────────────────────────────────────────────────────────────
    b.h1("12. Ожидаемые вопросы и развёрнутые ответы")

    def qa(question: str, answer: str):
        b.h3(f"Q: {question}")
        b.p(answer)

    qa(
        "Что вообще делает ваш проект, простыми словами?",
        "EventMind — это Telegram-бот, который собирает IT-мероприятия из "
        "разных источников, понимает их темы и уровень через LLM, и строит "
        "персональную ленту с объяснением «почему именно это». Плюс есть "
        "AI-помощник, который умеет составить план карьерного развития "
        "под конкретную цель.",
    )
    qa(
        "Почему гибридная рекомендательная система, а не одна нейросеть end-to-end?",
        "End-to-end решает только при больших данных, у нас прототип. Hybrid "
        "позволяет получать осмысленные результаты с первого пользователя: "
        "rule-компонент работает без истории, content — после первого "
        "embedding'а, bayesian/bandit/gnn — по мере накопления фидбэка. И "
        "каждый компонент объясним отдельно — у end-to-end такого нет.",
    )
    qa(
        "Как вы боретесь с галлюцинациями LLM?",
        "Четырьмя слоями: RAG (LLM видит только реальные данные из БД), "
        "Pydantic-валидация выходов, явный запрет на выдуманные факты в "
        "системных промптах, и обязательная ссылка на конкретный факт "
        "(тема, скилл, лайкнутое событие) в каждом объяснении.",
    )
    qa(
        "Чем LangGraph лучше LangChain Agents?",
        "Явный граф вместо «магической» цепочки. Каждая нода тестируема "
        "отдельно. Состояние типизировано через TypedDict. Conditional edges "
        "вместо if-в-промпте. Для multi-agent — гораздо чище.",
    )
    qa(
        "Что будет, если упадёт Groq?",
        "Три уровня fallback. Первый — автоматический переход на "
        "groq_fallback_model (8b вместо 70b). Второй — для /recommendations "
        "rule-компонент работает без LLM. Третий — /copilot возвращает "
        "«AI временно недоступен», но сессия сохраняется и при восстановлении "
        "сервиса всё продолжится.",
    )
    qa(
        "Как вы тестируете LLM-агентов?",
        "Тестируем не финальный текст (он недетерминирован), а: 1) что "
        "supervisor правильно классифицирует intent, 2) что нужные tools "
        "вызваны, 3) что pydantic-схема выходов прошла. Финальный текст — "
        "проверяется LLM-судьёй в scripts/llm_judge.py.",
    )
    qa(
        "Почему SQLite, а не Postgres?",
        "В dev — SQLite (быстрый старт, файл в репозитории). В проде — "
        "Postgres + pgvector. SQLAlchemy позволяет переключаться через "
        "DATABASE_URL. Миграция 9a1b2c3d4e5f включает pgvector только на PG; "
        "на SQLite embedding остаётся JSON в TEXT.",
    )
    qa(
        "Что нового по сравнению с обычным collaborative filtering?",
        "У нас CF — это только один из 9 компонент (gnn / LightGCN). Помимо "
        "него — content-based (cosine), Bayesian preference, contextual "
        "bandit (LinUCB), goal-aware (skill_gap), эвристики. CF в чистом "
        "виде страдает от cold-start; у нас он закрывается /bio + content.",
    )
    qa(
        "Какие метрики качества рекомендаций?",
        "Recall@k и nDCG@k через leave-one-out в scripts/eval_offline.py. "
        "Сравнивается rule / hybrid / hybrid+bandit. Плюс LLM-судья в "
        "scripts/llm_judge.py — он оценивает «осмысленность» объяснений.",
    )
    qa(
        "Что такое embedding и зачем он нужен?",
        "Embedding — это представление текста как точки в 384-мерном "
        "пространстве, в котором близкие по смыслу тексты лежат рядом. "
        "Используется для семантической близости (cosine), дедупликации, "
        "поиска по смыслу, MMR-диверсификации.",
    )
    qa(
        "Что такое RAG и где у вас применяется?",
        "Retrieval-Augmented Generation: сначала достаём релевантные "
        "документы, потом подсовываем LLM. У нас RAG в каждом AI-объяснении: "
        "hybrid отбирает top-N, LLM пишет только по ним — это убирает "
        "галлюцинации.",
    )
    qa(
        "Что такое function calling?",
        "Способность LLM сама решать, какую функцию вызвать. У нас 6 tools "
        "для Copilot: search_events, get_user_profile, "
        "get_interactions_summary, explain_event, recall_about_user, mark_saved.",
    )
    qa(
        "Что такое Thompson sampling и Bayesian-предпочтения?",
        "Beta-распределение моделирует «вероятность что тема понравится». "
        "На каждый фидбэк обновляем α (за положительный) и β (за отрицательный). "
        "Сэмплируем p — это и есть скор. Даёт автоматический баланс "
        "explore/exploit.",
    )
    qa(
        "Что такое LinUCB и чем отличается от Thompson sampling?",
        "LinUCB — контекстный бандит на линейной регрессии. Учитывает не "
        "одну тему, а 16-мерный вектор признаков пары (user, event). "
        "Thompson — глобальный по теме; LinUCB — локальный по контексту.",
    )
    qa(
        "Зачем MMR-диверсификация?",
        "Без неё топ-5 иногда состоит из 5 митапов про Kubernetes (все "
        "релевантные, но избыточные). MMR штрафует похожесть на уже "
        "выбранные — в топ попадают разные темы.",
    )
    qa(
        "Что такое long-term memory и чем она отличается от лайков?",
        "Лайки — это сырьё (телеметрия). Long-term memory — это активные "
        "выводы от первого лица: «хочет в ML», «не любит offline». Пишет "
        "LLM сам. Используется Copilot'ом — даёт ощущение «AI меня помнит».",
    )
    qa(
        "Можно ли это запустить у себя?",
        "Да. git clone + .env с BOT_TOKEN и GROQ_API_KEY + docker-compose "
        "up. Полное руководство — в README.md.",
    )
    qa(
        "Где у вас слабые места, если честно?",
        "GNN на маленьком датасете шумит — выключен по умолчанию. На SQLite "
        "редкие «database is locked» при параллельном ingestion (на PG не "
        "воспроизводится). MMR-λ подобран эмпирически. LLM-судья — пока "
        "однократный.",
    )
    b.empty()

    # ── 13 ───────────────────────────────────────────────────────────────
    b.h1("13. Шпаргалка: команды для копи-паста")
    b.h3("Запуск всего")
    b.bullet("git clone … && cd eventmind")
    b.bullet("cp .env.example .env  # вписать BOT_TOKEN и GROQ_API_KEY")
    b.bullet("docker-compose up -d")
    b.bullet("docker logs eventmind-api-1 -f")
    b.bullet("alembic upgrade head  # если первый запуск")
    b.bullet("В Telegram: /start")
    b.h3("Ingestion на демо")
    b.bullet("curl -X POST 'localhost:8000/ingestion/load-habr?limit=5'")
    b.bullet("curl -X POST 'localhost:8000/ingestion/load-rss?limit_per_feed=5'")
    b.bullet("curl 'localhost:8000/events/search?q=python'")
    b.h3("Проверки")
    b.bullet("curl localhost:8000/health")
    b.bullet("curl localhost:8000/recommendations/<telegram_id>")
    b.bullet("curl 'localhost:8000/analytics/trending?days=7'")
    b.h3("Тесты")
    b.bullet("pytest -q")
    b.bullet("ruff check .")
    b.h3("Сценарии в боте")
    b.bullet("/start → тур → 🚀 Начать настройку")
    b.bullet("/recommend → 👍 / 👎 / ⭐ / Похожие / Следующее")
    b.bullet("❓ Почему → короткий ответ в поп-апе; 📖 Подробнее → разбор")
    b.bullet("/copilot хочу за полгода в middle backend")
    b.bullet("/bio backend 3 года, хочу в ML")
    b.bullet("/edit — перенастроить темы/формат/город")
    b.bullet("/undo — откатить последний фидбэк")
    b.h3("Полезные curl для админки")
    b.bullet("curl localhost:8000/admin/stats")
    b.bullet("curl localhost:8000/subscriptions/users")
    b.bullet("curl localhost:8000/vocabulary/topics")
    b.empty()

    # ── 14 ───────────────────────────────────────────────────────────────
    b.h1("14. Финальные советы по защите")
    b.bullet("Открыть docs/diagrams/architecture.png и держать на втором мониторе — это якорь")
    b.bullet("Если вопрос «как вы это решаете» — сначала назвать парадигму (content / collaborative / contextual), потом конкретный компонент")
    b.bullet("Никогда не скрывать слабые места: «GNN выключен на малом датасете» — это плюс, а не минус")
    b.bullet("Если LLM отвечает странно на демо — спокойно: «У нас есть fallback на rule-only, я переключусь»")
    b.bullet("Подсветить число тестов: 180, и что они проверяют не текст LLM, а граф/tools/схемы")


# ════════════════════════════════════════════════════════════════════════
#  Отчёт по курсовой работе
# ════════════════════════════════════════════════════════════════════════


def build_thesis(b: Builder, *, default_text: str, list_style: str) -> None:
    """Полный технический отчёт. Здесь b.h1 = HSE Title 1, b.h2 = HSE Title 2,
    b.h3 = HSE Title 3, b.p = HSE Default Text (через override style)."""

    # Так как HSE-стили нужны не для всех частей, переопределяем methods Builder'а локально.
    # На уровне функции мы используем переданные стили.

    def p(text: str) -> None:
        b.doc.add_paragraph(text, style=default_text)

    def li(text: str) -> None:
        b.doc.add_paragraph(text, style=list_style)

    # ── Титульный блок ───────────────────────────────────────────────────
    p("")
    p("")
    p("Министерство науки и высшего образования Российской Федерации")
    p("Национальный исследовательский университет «Высшая школа экономики»")
    p("")
    p("Пермский кампус")
    p("Факультет экономики, менеджмента и бизнес-информатики")
    p("")
    p("")
    p("Курсовая работа")
    p("«EventMind — система агрегации IT-мероприятий и персонализированных AI-рекомендаций»")
    p("")
    p("")
    p("Пермь, 2026 год")

    # ── Аннотация ────────────────────────────────────────────────────────
    b.h1("Аннотация")
    p(
        "В работе рассматривается проектирование и реализация программной "
        "системы для агрегации IT-мероприятий и персонализированных AI-"
        "рекомендаций. Система собирает анонсы из шести независимых "
        "источников (Habr, RSS/Atom, KudaGo, Lu.ma, Meetup, Telegram-каналы), "
        "приводит их к единой структуре через LLM-нормализацию с Pydantic-"
        "валидацией и хранит в реляционной БД с поддержкой векторного типа "
        "(pgvector). Семантическая дедупликация через cosine ≥ 0.92 "
        "устраняет парафразы между источниками."
    )
    p(
        "Рекомендательная подсистема построена как гибридная: ранжирование "
        "выполняется суммой девяти компонент. Помимо классических content-"
        "based (cosine) и rule-based, реализованы Bayesian preference "
        "(Thompson sampling по темам с temporal decay), контекстный бандит "
        "(LinUCB на 16-мерном векторе пары пользователь–событие) и "
        "коллаборативная компонента (LightGCN на двудольном графе user↔event). "
        "Поверх — MMR-диверсификация (λ=0.7) и LLM-объяснение выдачи. "
        "Контекстная и коллаборативная парадигмы реализованы в одном "
        "ранжировании одновременно."
    )
    p(
        "Спроектирована схема базы данных из 11 таблиц с разделением "
        "«пользователь — событие — взаимодействие — служебные подсистемы». "
        "Реализован пайплайн ingestion с LLM-нормализатором на LangGraph; "
        "REST API на FastAPI с 11 роутерами; Telegram-бот на aiogram 3 с "
        "онбординг-туром, лентой рекомендаций, объяснимостью двух уровней "
        "и мульти-туровым AI-копилотом. Copilot собран как граф LangGraph "
        "Supervisor-Worker с пятью специалистами и пятью function-calling "
        "tools, с поддерживаемой через таблицу copilot_sessions памятью "
        "диалога и подключением к долговременной памяти пользователя "
        "(mem0-стиль). Реализован планировщик фоновых задач (APScheduler): "
        "AI-дайджест, периодический ingestion и уплотнение памяти."
    )

    # ── Оглавление-плейсхолдер ───────────────────────────────────────────
    b.h1("Оглавление")
    p("(оглавление формируется средствами Word: Ссылки → Оглавление)")

    # ── Введение ─────────────────────────────────────────────────────────
    b.h1("Введение")
    p(
        "Растущее количество IT-мероприятий — конференций, митапов, "
        "вебинаров, школ — порождает обратную проблему: анонсы рассредоточены "
        "по десяткам каналов, форматы публикации различаются, ручной поиск "
        "релевантного отнимает значимое время. Цель курсовой работы — "
        "спроектировать и реализовать программную систему, которая решает "
        "эту проблему сквозным образом: автоматически агрегирует события "
        "из множества источников, приводит их к единой структуре через "
        "AI-обогащение, строит персонализированную ленту с прозрачным "
        "объяснением «почему именно это» и поддерживает мульти-туровый "
        "диалог с AI-помощником по карьерным задачам пользователя."
    )
    p(
        "Задачи работы: 1) проанализировать предметную область и обзор "
        "существующих решений, выделить требования; 2) спроектировать "
        "архитектуру системы из независимых процессов и схему БД; 3) "
        "реализовать пайплайн ingestion с шестью адаптерами и AI-"
        "нормализатором; 4) построить гибридную рекомендательную "
        "подсистему, объединяющую content-based, контекстный бандит, "
        "коллаборативную и Bayesian-парадигмы; 5) реализовать AI-копилота "
        "как multi-agent граф LangGraph с долговременной памятью; 6) "
        "обеспечить REST API, Telegram-бота и контейнеризованное "
        "развёртывание; 7) провести функциональное тестирование (180 "
        "автотестов), офлайн-оценку рекомендера (Recall@k, nDCG@k) и "
        "LLM-оценку качества выдачи."
    )
    p(
        "Научно-практическая новизна состоит в одновременной реализации в "
        "одном ранжировании двух разнородных парадигм — контекстного "
        "бандита (LinUCB) и коллаборативной фильтрации на графе "
        "(LightGCN), — каждая из которых выписана с нуля на чистом numpy "
        "и интегрирована в hybrid через единый интерфейс компоненты. "
        "В UX-плане новизна заключается в двухуровневой объяснимости (поп-"
        "ап и развёрнутый разбор с counterfactual) и в mem0-стилевой "
        "долговременной памяти пользователя, подключённой к AI-копилоту."
    )

    # ════════════════════════════════════════════════════════════════════
    # Глава 1
    # ════════════════════════════════════════════════════════════════════
    b.h1("Глава 1. Анализ предметной области")

    b.h2("1.1. Описание предметной области и постановка проблемы")
    p(
        "Предметная область — рынок IT-мероприятий и потребление контента "
        "профессиональным сообществом. Под IT-мероприятиями понимаются "
        "конференции, митапы, вебинары, воркшопы, школы, хакатоны и "
        "связанные форматы, проводимые офлайн, онлайн или в гибридном "
        "режиме. Источники анонсов — официальные сайты, корпоративные "
        "блоги, новостные RSS-ленты, специализированные платформы (KudaGo, "
        "Lu.ma, Meetup), Telegram-каналы."
    )
    p(
        "Проблема пользователя: высокая фрагментация. Один и тот же митап "
        "может одновременно публиковаться в трёх местах под разными "
        "заголовками; форматы описания (поля «уровень», «темы», «стек») "
        "не нормализованы; одни источники структурированы (Meetup API), "
        "другие — нет (Telegram-каналы). В результате задача «найти "
        "релевантное мне» сводится к ручному просмотру множества каналов "
        "и решается крайне неэффективно."
    )
    p(
        "Проблема рекомендера: даже после агрегации стандартного "
        "коллаборативного фильтра недостаточно. У пользователя может быть "
        "явно сформулированная карьерная цель («перейти из backend в ML "
        "за 6 месяцев»), которая должна учитываться при ранжировании, и "
        "запрос «объясни, почему ты порекомендовал X» должен получать "
        "осмысленный ответ, а не «вы похожи на пользователя Y, который "
        "это лайкнул»."
    )

    b.h2("1.2. Обзор существующих решений")
    p(
        "Прямых конкурентов на российском рынке немного. Meetup.com — "
        "глобальная платформа с собственной лентой, однако охватывает "
        "ограниченный круг русскоязычных мероприятий и не делает AI-"
        "объяснений. Timepad/Eventbrite — каталоги без персонализации. "
        "KudaGo — городской афишер с категорией «IT», но без модели "
        "пользователя. Telegram-каналы вроде @iteventsru — высокопокрытый "
        "ручной канал без рекомендаций. Корпоративные ленты (Yandex/VK) — "
        "узкие по охвату."
    )
    p(
        "Из академических наработок наиболее близки гибридные "
        "рекомендатели для образовательного контента (Coursera, edX), "
        "однако они не работают с потоком ежедневных коротких событий "
        "(митапов) и не интегрируют LLM-объяснения. Известные индустриальные "
        "продукты с AI-объяснениями (Spotify Discover, Netflix Why-Watching) "
        "не публикуют детали алгоритмов."
    )

    b.h2("1.3. Рекомендательные системы и подходы к персонализации")

    b.h3("1.3.1. Контентно-ориентированные рекомендации")
    p(
        "Сравнивают «отпечаток» пользователя (выбранные темы, прошлые "
        "лайки, профильный embedding) с описаниями событий. Преимущество — "
        "холодный старт хорошо закрывается даже первым лайком. Недостаток — "
        "склонность к «эхо-камере» (рекомендуется похожее, отсутствует "
        "разнообразие)."
    )

    b.h3("1.3.2. Коллаборативная фильтрация")
    p(
        "Использует поведенческие данные «других похожих пользователей». "
        "Современная реализация — графовые модели (LightGCN, NGCF). "
        "Преимущество — высокое качество при плотном графе. Недостаток — "
        "холодный старт и неинтерпретируемость («так лайкали другие»)."
    )

    b.h3("1.3.3. Гибридные рекомендательные системы")
    p(
        "Объединяют content-based, collaborative, контекстные и rule-based "
        "сигналы в один score. Преимущества каждого подхода складываются, "
        "недостатки взаимно компенсируются. В работе реализован hybrid из "
        "9 компонент с прозрачной разбивкой вклада каждой."
    )

    b.h3("1.3.4. Семантический анализ и embeddings")
    p(
        "Семантическая близость через cosine-similarity на эмбеддингах "
        "решает задачу «понять смысл, а не только слова». Используется "
        "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2 (384 "
        "dim), поддерживающий русский. Embeddings применяются в трёх "
        "местах: рекомендатель (cosine-компонент), retrieval для Copilot "
        "(top-k по семантике), дедупликация ingestion (cosine ≥ 0.92)."
    )

    b.h3("1.3.5. Использование обратной связи пользователя")
    p(
        "Реализованы три формы фидбэка: лайк, дизлайк, save. Каждый "
        "обновляет четыре подсистемы синхронно в одной транзакции: "
        "topic_weights, alpha/beta в Bayesian-модели, матрицу A и вектор "
        "b в контекстном бандите, заметки long-term памяти. Команда /undo "
        "откатывает все четыре подсистемы."
    )

    b.h2("1.4. Источники данных и способы агрегации")
    b.h3("1.4.1. Источники данных о мероприятиях")
    p(
        "Шесть категорий источников: HTML-парсинг (Habr-анонсы), RSS/Atom "
        "(произвольный список feed'ов), публичный JSON API (KudaGo), "
        "ICS-фиды (Lu.ma), GraphQL API (Meetup), web-preview Telegram-"
        "каналов (страница t.me/s/<имя>)."
    )
    b.h3("1.4.2. Проблемы агрегации данных")
    p(
        "Главные сложности: разнородность форматов, дубли между "
        "источниками, шумные не-IT записи, неполнота полей (часто нет "
        "уровня / стека). Решение — AI-нормализация LLM с pydantic-"
        "схемой выхода, семантическая дедупликация, фильтрация по topics."
    )
    b.h3("1.4.3. Нормализация данных")
    p(
        "EventNormalizerAgent — одна нода LangGraph с одним LLM-вызовом. "
        "Возвращает JSON с фиксированной структурой; pydantic проверяет "
        "типы. Если topics == [] — событие признаётся не-IT, статус "
        "raw_event=non_it. Заполняются quality_score и hype_score в "
        "шкале 1–10."
    )

    b.h2("1.5. Использование искусственного интеллекта в рекомендательных системах")
    p(
        "AI применяется в проекте на четырёх уровнях: 1) нормализация "
        "сырых событий через LLM с pydantic-валидацией, 2) объяснения "
        "ранжирования через короткий и развёрнутый LLM-ответы, 3) мульти-"
        "агентный AI-копилот для карьерных задач, 4) долговременная "
        "память пользователя — LLM-агент пишет активные заметки от "
        "первого лица."
    )

    b.h2("1.6. Требования к разрабатываемой системе")
    b.h3("1.6.1. Функциональные требования")
    li("Сбор событий из шести источников")
    li("AI-нормализация с pydantic-валидацией")
    li("Семантическая дедупликация")
    li("Регистрация и настройка профиля пользователя")
    li("Персонализированные рекомендации с прозрачным объяснением")
    li("Поиск событий по тексту и семантике")
    li("Сохранение событий и обратная связь (лайк/дизлайк/save/undo)")
    li("AI-копилот с мульти-туровым диалогом и памятью")
    li("Подписка на AI-дайджест")
    b.h3("1.6.2. Требования к рекомендательной подсистеме")
    li("Гибридный подход из не менее чем 5 разнородных компонент")
    li("Контекстная и коллаборативная парадигмы — каждая независимо реализована")
    li("Объяснимость каждой выдачи")
    li("Reversible feedback (откат)")
    li("Холодный старт через /bio")
    b.h3("1.6.3. Требования к подсистеме обработки данных")
    li("Единый интерфейс источника")
    li("AI-нормализация через LLM")
    li("Семантическая дедупликация cosine ≥ 0.92")
    li("Атомарность сохранения и status-pipeline (raw → normalized / non_it)")
    b.h3("1.6.4. Нефункциональные требования")
    li("Время ответа /recommendations ≤ 1 c при 200 событиях в базе")
    li("Безотказность: сбой LLM не должен приводить к пустой выдаче")
    li("Изолируемость подсистем: сбой long-term memory не ломает транзакцию фидбэка")
    li("Контейнеризация и healthcheck-ориентированный старт")
    b.h3("1.6.5. Требования к пользовательскому интерфейсу")
    li("Telegram-бот как единственная точка входа")
    li("Reply-меню с 4 главными кнопками")
    li("Inline-кнопки в карточках с обратной связью и объяснениями")
    li("Объяснимость двух уровней: короткое в поп-апе и развёрнутое отдельным сообщением")
    b.h3("1.6.6. Требования к архитектуре системы")
    li("Три независимых процесса: API, бот, scheduler")
    li("Общая БД, REST-общение между процессами")
    li("Миграции схемы под контролем версий")
    li("Возможность переключения SQLite ↔ PostgreSQL + pgvector через переменную окружения")

    b.h2("1.7. Сценарии использования (use cases)")
    b.h3("1.7.1. Регистрация и настройка профиля пользователя")
    p(
        "Пользователь стартует бота командой /start, видит приветствие и "
        "короткий 4-экранный тур по командам. На последнем шаге тура "
        "появляется инлайн-кнопка «🚀 Начать настройку»: выбор тем (можно "
        "несколько), формата (online/offline/hybrid/any), города. "
        "Профиль сохраняется через POST /users/register. Альтернативный "
        "путь — /bio с описанием себя в свободной форме (cold-start через "
        "LLM-извлечение скилл-профиля). Тур можно пропустить — после "
        "пропуска бот всё равно показывает кнопку запуска настройки."
    )
    b.h3("1.7.2. Получение рекомендаций")
    p(
        "Пользователь жмёт кнопку «🎯 Рекомендации» reply-меню или "
        "команду /recommend. Бэкенд считает hybrid_score по 9 компонентам, "
        "делает MMR-диверсификацию и возвращает упорядоченный список. Бот "
        "рисует первую карточку и инлайн-кнопки: ❓ Почему, 📖 Подробнее, "
        "👍, 👎, ⭐, Похожие, Следующее. Кнопки ❓ Почему и 📖 Подробнее — "
        "независимые: первая открывает поп-ап с коротким LLM-объяснением, "
        "вторая шлёт отдельное сообщение с развёрнутым разбором, "
        "counterfactual'ом и числовой разбивкой по компонентам."
    )
    b.h3("1.7.3. Поиск и фильтрация событий")
    p(
        "Пользователь жмёт кнопку «🔍 Поиск» и вводит запрос. Один поиск "
        "делает два прохода одновременно: keyword (до 3 событий) и "
        "семантический по эмбеддингам (до 5 событий с дедупом по id из "
        "keyword-блока). Запросом может быть как точная фраза («python "
        "митап»), так и формулировка цели («хочу научиться деплоить "
        "микросервисы») — combined_search распознаёт goal-intent и "
        "перекидывает в Copilot."
    )
    b.h3("1.7.4. Подписка на рекомендации и уведомления")
    p(
        "Пользователь жмёт /subscribe или включает подписку в профиле. "
        "Scheduler-процесс раз в DIGEST_INTERVAL_MINUTES (по умолчанию 1440) "
        "запрашивает по каждому подписчику его топ-карточку AI-рекомендации "
        "и отправляет через Telegram Bot API. Для разработки интервал "
        "ставится в 10 минут с DIGEST_RUN_ON_STARTUP=true."
    )

    b.h2("1.8. Бизнес-процессы")
    b.h3("1.8.1. Процесс сбора данных о событиях")
    p(
        "Scheduler по интервалу INGEST_INTERVAL_HOURS дёргает три "
        "endpoint'а API: /ingestion/load-habr, /ingestion/load-rss, "
        "/ingestion/load-tg. Каждый — обёртка над соответствующим "
        "адаптером в app/ingestion/sources. Адаптер возвращает список "
        "raw_event-объектов; сервис записывает их в таблицу raw_events со "
        "status=raw."
    )
    b.h3("1.8.2. Процесс обработки и нормализации данных")
    p(
        "После записи raw_events запускается EventNormalizerAgent: для "
        "каждой записи — один LLM-вызов с запросом «извлеки title, format, "
        "city, level, date, topics, tech_stack, quality_score, hype_score "
        "из этого сырого описания». Результат валидируется pydantic-схемой. "
        "При topics == [] статус меняется на non_it. Иначе — семантическая "
        "дедупликация (cosine ≥ 0.92 за 90 дней), и при отсутствии дубля "
        "запись создаётся в events с embedding'ом и status=normalized."
    )
    b.h3("1.8.3. Процесс формирования пользовательской модели")
    p(
        "При каждом действии пользователя обновляются: topic_weights в "
        "users (rule-сигнал), alpha/beta в user_topic_stats (Bayesian), "
        "матрица A и вектор b в user_bandit_states (LinUCB), embedding "
        "пользователя в users.embedding (cosine). Опционально — заметка "
        "long-term памяти в user_memories."
    )
    b.h3("1.8.4. Процесс формирования рекомендаций")
    p(
        "Запрос /recommendations/{tid}: достать все события → прогреть "
        "кэши embedding'ов → для каждого события подсчитать 9 компонент "
        "→ суммировать → отсортировать → MMR-rerank → вернуть JSON."
    )
    b.h3("1.8.5. Процесс взаимодействия пользователя с системой")
    p(
        "Пользователь общается с системой исключительно через Telegram-"
        "бот. Бот — тонкий клиент, бизнес-логики в нём минимум, всё "
        "выносится в REST API. Это даёт две выгоды: 1) одинаковый бэк "
        "может обслуживать веб-клиент / мобильное приложение в будущем; "
        "2) тестируется отдельно — pytest на API без поднятия Telegram."
    )

    # ════════════════════════════════════════════════════════════════════
    # Глава 2
    # ════════════════════════════════════════════════════════════════════
    b.h1("Глава 2. Проектирование системы")

    b.h2("2.1. Проектирование архитектуры системы")
    b.h3("2.1.1. Общая архитектура системы")
    p(
        "Принят паттерн «три независимых процесса вокруг общей БД»: api "
        "(REST), bot (Telegram long-polling), scheduler (фоновые джобы). "
        "Каждый процесс — отдельный Docker-контейнер с собственным "
        "жизненным циклом. Бот и scheduler ждут healthcheck'а api перед "
        "стартом. Это даёт следующие гарантии: бэкенд можно перезапустить "
        "без перезапуска бота (теряются inflight-запросы, но не сессии); "
        "scheduler можно отключить полностью (через INGEST_ENABLED=false), "
        "не задевая API; добавление новой клиентской подсистемы (например, "
        "веб-фронта) не требует трогать существующие процессы."
    )
    b.h3("2.1.2. Архитектура backend-приложения")
    p(
        "Внутри API применяется слой роутеров (FastAPI APIRouter), слой "
        "сервисов (бизнес-логика, app/api/services) и слой моделей "
        "(SQLAlchemy, app/db/models). Роутер только парсит вход через "
        "Pydantic-схему и дёргает сервис. Сервисный слой принимает Session, "
        "оперирует моделями, возвращает dict/Pydantic. Это даёт чистое "
        "тестирование (pytest вызывает сервис напрямую) и возможность "
        "переиспользовать сервисы из scheduler-процесса."
    )
    b.h3("2.1.3. Архитектура Telegram-бота")
    p(
        "aiogram 3 организован вокруг router'ов: 6 router'ов "
        "(start, recommendations, profile, search, copilot, subscriptions) "
        "регистрируются в нужном порядке (start первым — он перехватывает "
        "deep-link; copilot последним — он catch-all для продолжения "
        "сессии). Все вызовы REST API делаются через единый HTTP-клиент "
        "(app/bot/services/api_client.py) с явными таймаутами и фолбэками. "
        "Состояние мульти-турового Copilot'а живёт в RAM бота "
        "(user_id → {session_id, last_activity}), а сама история — на "
        "бэке в copilot_sessions."
    )
    b.h3("2.1.4. Архитектура AI-модулей и агентов")
    p(
        "AI-логика разделена на две папки: app/agents/recommendation/ — "
        "только EventNormalizerAgent (одна нода + пользователь его не "
        "видит) и app/agents/copilot/ — Supervisor-Worker граф LangGraph "
        "с пятью специалистами и пятью function-calling tools. Каждый "
        "специалист — отдельный модуль app/agents/copilot/specialists/, "
        "tools — общий app/agents/copilot/tools.py. Логика LLM-вызова и "
        "fallback'а на резервную модель централизована в "
        "app/agents/recommendation/llm.py."
    )

    b.h2("2.2. Проектирование базы данных")
    b.h3("2.2.1. Основные сущности системы")
    li("users — профиль пользователя (telegram_id, preferred_format, city, topic_weights JSON, embedding)")
    li("topics — справочник тем")
    li("events — нормализованные события")
    li("event_topics — many-to-many")
    li("interactions — лайки/дизлайки/save")
    li("user_topics — many-to-many для выбранных пользователем тем")
    li("user_topic_stats — Bayesian alpha/beta по темам")
    li("user_bandit_states — A, b для LinUCB")
    li("user_skill_profiles — карьерный профиль")
    li("user_memories — заметки long-term памяти")
    li("copilot_sessions — история мульти-туровых диалогов")
    li("raw_events — сырые записи ingestion (status=raw/normalized/non_it)")
    b.h3("2.2.2. Связи между сущностями")
    p(
        "users 1-к-N interactions N-к-1 events; events N-к-N topics; users "
        "N-к-N topics; users 1-к-1 user_skill_profiles; users 1-к-N "
        "user_topic_stats (по теме); users 1-к-1 user_bandit_states; users "
        "1-к-N user_memories; users 1-к-N copilot_sessions."
    )
    b.h3("2.2.3. Структура хранения событий")
    p(
        "Двухслойное хранение: raw_events (полный сырой payload + статус) "
        "и events (нормализованные структурированные поля). Это даёт "
        "трассируемость («покажи мне сырой текст для события 42») и "
        "возможность переобработать неудачные ingestion'ы без повторной "
        "выгрузки из источника."
    )
    b.h3("2.2.4. Структура хранения пользовательских данных")
    p(
        "users — «горячее ядро» (часто обновляемое: embedding, "
        "topic_weights). Sidecar-таблицы (user_topic_stats, "
        "user_bandit_states, user_skill_profiles, user_memories) хранят "
        "узкоспециализированное состояние подсистем. Это разделение даёт "
        "независимые миграции и возможность выключать отдельные "
        "подсистемы через feature flag'и без миграции схемы."
    )

    b.h2("2.3. Проектирование пайплайна обработки событий")
    b.h3("2.3.1. Сбор данных из внешних источников")
    p(
        "Единый интерфейс EventSource: метод fetch() возвращает список "
        "RawEvent-объектов. Каждый адаптер реализует этот интерфейс "
        "своим способом. Конкретные адаптеры включают/выключаются "
        "через переменные окружения (RSS_FEEDS, MEETUP_TOKEN, "
        "LUMA_CALENDARS, TG_INGEST_CHANNELS)."
    )
    b.h3("2.3.2. Нормализация данных")
    p(
        "Один LLM-вызов на сырой event с пакетом полей в выходной "
        "Pydantic-схеме. Если схема не проходит — событие маркируется "
        "статусом failed, в логи пишется причина. Это даёт явный feedback "
        "на нестабильные источники: «адаптер X стал отдавать поля в "
        "новом формате»."
    )
    b.h3("2.3.3. Обнаружение дубликатов")
    p(
        "Перед сохранением нормализованного события считаем cosine между "
        "его embedding'ом и embedding'ами событий за последние 90 дней. "
        "При значении ≥ DEDUP_THRESHOLD (0.92) сохранение пропускается. "
        "Этот выбор порога подобран эмпирически: 0.95 пропускает парафразы, "
        "0.90 ловит ложные совпадения."
    )
    b.h3("2.3.4. Обогащение данных с помощью AI")
    p(
        "Помимо извлечения структуры, LLM-нормализатор выставляет "
        "quality_score (1–10, оценка содержательности) и hype_score (1–10, "
        "громкость). Эти два скора напрямую участвуют в hybrid-ранжировании "
        "как отдельные компоненты."
    )

    b.h2("2.4. Проектирование рекомендательной системы")
    b.h3("2.4.1. Формирование модели пользователя")
    p(
        "Модель пользователя — шесть слоёв: базовый профиль (users), "
        "выбранные темы (user_topics), Bayesian-статистика "
        "(user_topic_stats), skill-профиль (user_skill_profiles), "
        "состояние бандита (user_bandit_states), долговременная память "
        "(user_memories). Каждый слой обновляется по своему триггеру и "
        "используется конкретной компонентой hybrid-скоринга."
    )
    b.h3("2.4.2. Rule-based рекомендации")
    p(
        "Правила покрывают «холодный» сценарий, когда нет ни истории, "
        "ни эмбеддингов: совпадение тем, формата, города. Вес — 0.5, "
        "наименьший в hybrid'е (rule намеренно «придерживается», уступая "
        "более тонким сигналам), но компонент гарантирует осмысленную "
        "выдачу даже при полном отказе AI-стека."
    )
    b.h3("2.4.3. Семантический поиск и embeddings")
    p(
        "Все сравнения «по смыслу» делаются через cosine на эмбеддингах "
        "MiniLM (384 dim). Embedding пользователя строится из тем, "
        "форматов, истории взаимодействий (build_rich_user_embedding); "
        "embedding события — из title + description. Оба кэшируются в "
        "соответствующих таблицах. На PG используется тип vector и "
        "оператор <=> для быстрого top-k через индексы."
    )
    b.h3("2.4.4. Гибридный алгоритм рекомендаций")
    p(
        "Реализован hybrid из девяти компонент: rule, cosine, bayesian, "
        "quality, hype, freshness, skill_gap, bandit, gnn. Каждая возвращает "
        "float (вклад с учётом веса) или 0.0 при сбое. Сбой одной компоненты "
        "не ломает остальные. Итог — сумма; сортировка по убыванию; "
        "пост-процессинг MMR-диверсификации (λ=0.7)."
    )
    b.h3("2.4.5. Использование обратной связи")
    p(
        "Каждый лайк/дизлайк/save обновляет ЧЕТЫРЕ подсистемы синхронно: "
        "topic_weights, alpha/beta, A/b LinUCB, заметки памяти. /undo — "
        "симметричный rollback с direction=-1. Это нетривиально: в "
        "большинстве рекомендеров «обратной операции» для бандита и "
        "Bayesian-модели нет."
    )

    b.h2("2.5. Выбор технологий и средств реализации")
    b.h3("2.5.1. Язык программирования и backend")
    p("Python 3.12 + FastAPI 0.110 + Uvicorn. Причины — экосистема AI/ML, нативная асинхронность, скорость разработки.")
    b.h3("2.5.2. Telegram Bot Framework")
    p("aiogram 3 — современный Python-фреймворк с router'ами и FSM. Альтернатива (python-telegram-bot) менее удобна для multi-router архитектуры.")
    b.h3("2.5.3. ORM и база данных")
    p("SQLAlchemy 2.0 + Alembic. SQLite для разработки (файл в репо), PostgreSQL + pgvector для продакшна. Переключение через DATABASE_URL.")
    b.h3("2.5.4. AI-инструменты и языковые модели")
    p(
        "LangGraph 0.2 для multi-agent. LLM-цепочка _LLMChain: первичный "
        "провайдер Google Gemini через langchain-google-genai с REST-"
        "транспортом (REST устойчивее gRPC к местным DNS-фильтрам), модель "
        "выбирается автопробой на старте API. Fallback'и — Groq llama-3.3-70b "
        "и llama-3.1-8b. Поверх — circuit-breaker и per-provider cooldown. "
        "sentence-transformers 2.7 для embedding'ов (многоязычный MiniLM-L12-v2)."
    )
    b.h3("2.5.5. Средства контейнеризации и развертывания")
    p("Docker + docker-compose: api, bot, scheduler — три сервиса с общей БД. Healthcheck перед стартом зависимых.")

    # ════════════════════════════════════════════════════════════════════
    # Глава 3
    # ════════════════════════════════════════════════════════════════════
    b.h1("Глава 3. Реализация системы")

    b.h2("3.1. Реализация базы данных и слоя доступа")
    b.h3("3.1.1. Модели SQLAlchemy")
    p(
        "11 моделей в app/db/models/: User, Event, RawEvent, Interaction, "
        "Topic, CopilotSession, UserBanditState, UserSkillProfile, "
        "UserTopicStat, UserMemory + ассоциации event_topics, user_topics. "
        "Используются relationship() с lazy='joined' для часто запрашиваемых "
        "связей. Embedding-поля хранятся как Text (JSON-строка) на SQLite и "
        "как vector(384) на PostgreSQL — это разрешает диалект-специфичный "
        "тип через TypeDecorator."
    )
    b.h3("3.1.2. Миграции Alembic")
    p(
        "Десять миграций: initial_schema (6af520bb31e0), enriched_fields "
        "(b3c4d5e6f7a8), summary_embedding (7b1c9d2e3a4f), user_skill_profile "
        "(d6e7f8901234), user_topic_stats (c5d6e7f89012), copilot_sessions "
        "(e7f890123456), bandit_state (f8901234abcd), user_memories "
        "(abcd1234ef56), pgvector (9a1b2c3d4e5f). Каждая независимо "
        "ревёрсится через downgrade()."
    )
    b.h3("3.1.3. Обеспечение конкурентного доступа к SQLite")
    p(
        "Подключён глобальный event-listener на connect: PRAGMA journal_mode=WAL "
        "(читатели не блокируют писателей), busy_timeout=10000 (10 с ожидания "
        "вместо мгновенного fail), synchronous=NORMAL (быстрее на записи, "
        "безопасно с WAL). Кроме того, все сервисные операции, способные "
        "конкурировать с ingestion'ом (например, refresh_user_embedding), "
        "оборачиваются в try/except с явным db.rollback() — это критично, "
        "так как непровёрнутый flush оставляет сессию в pending-rollback "
        "состоянии и ломает последующие запросы."
    )

    b.h2("3.2. Реализация REST API")
    b.h3("3.2.1. Структура роутеров")
    p(
        "11 роутеров: users, events, recommendations, agent_recommendations, "
        "copilot, ingestion, analytics, subscriptions, vocabulary, admin, "
        "плюс служебный composite /health. Каждый — APIRouter с prefix и "
        "tags. Подключение в app/api/main.py через include_router."
    )
    b.h3("3.2.2. Сервисный слой")
    p(
        "В app/api/services лежат: recommendation_service, "
        "ai_recommendation_service, search_service, subscription_service, "
        "user_service, event_service, ingestion_service. Роутер только "
        "передаёт Session + payload, сервис делает всю бизнес-логику. Это "
        "позволяет тестировать сервисы напрямую через pytest без поднятия "
        "сервера."
    )
    b.h3("3.2.3. Жизненный цикл приложения и прогрев модели embeddings")
    p(
        "На старте API инициализируется sentence-transformers модель "
        "(операция тяжёлая, ~1.5 с). Без прогрева первый запрос "
        "/recommendations занимал ~16 с — теперь стабильно ~0.4 с после "
        "первого запроса. Кэширование событий через ensure_event_embeddings "
        "(пакетная инициализация) — батч-encode, а не N отдельных вызовов."
    )
    b.h3("3.2.4. Валидация и обработка ошибок")
    p(
        "Pydantic v2 schemas для входа/выхода. На сервисном уровне — "
        "HTTPException с кодами 404 (не найдено), 400 (некорректный action), "
        "и явные fallback-ответы при сбое LLM-цепочки. Все exception'ы в "
        "/copilot и /agent-recommendations логируются через logger.exception "
        "(добавлено 29 мая) — раньше generic-сообщение «временно недоступен» "
        "приходило без следа в логах, и причину нельзя было выяснить."
    )

    b.h2("3.3. Реализация Telegram-бота")
    b.h3("3.3.1. Точка входа и инициализация")
    p(
        "app/bot/main.py создаёт Bot и Dispatcher, регистрирует 6 router'ов "
        "в строгом порядке (start первым, copilot последним — он catch-all), "
        "и стартует long-polling. Порядок важен: точные F.text-фильтры "
        "должны срабатывать раньше catch-all с pending-state."
    )
    b.h3("3.3.2. Обработчики команд и сообщений")
    p(
        "Команды: /start (с deep-linking для шеринга событий), /tour, "
        "/recommend, /search, /subscribe, /unsubscribe, /profile, /edit, "
        "/bio, /undo, /copilot, /trending, /help. Reply-кнопки: «🎯 "
        "Рекомендации», «🔍 Поиск», «👤 Профиль», «⚙️ Ещё». На /start "
        "теперь принудительно сбрасывается старая reply-клавиатура через "
        "ReplyKeyboardRemove — иначе у пользователей старой версии "
        "оставалась залипшая кнопка «Начать настройку»."
    )
    b.h3("3.3.3. Inline-клавиатуры")
    p(
        "В app/bot/keyboards/inline.py — фабрики клавиатур: tour_keyboard, "
        "topics_keyboard, format_keyboard, city_keyboard, "
        "recommendation_keyboard (с кнопками ❓ Почему / 📖 Подробнее / "
        "👍 / 👎 / ⭐ / Похожие / Следующее), profile_actions_keyboard "
        "(включая 📬 AI-дайджест: ВКЛ/ВЫКЛ), more_menu_keyboard."
    )
    b.h3("3.3.4. Сервисный слой для взаимодействия с API")
    p(
        "EventMindAPIClient — единая обёртка над httpx.AsyncClient. "
        "Глобальный таймаут — 30 с для большинства методов, 60 с для "
        "Copilot. Каждый метод ловит httpx.HTTPError и возвращает безопасное "
        "значение по умолчанию (пустой dict, пустой список) — это "
        "значит, что бот никогда не падает Telegram-исключением при сбое "
        "API, всегда показывает «временно недоступно» в UI."
    )

    b.h2("3.4. Реализация подсистемы ингеста")
    b.h3("3.4.1. Унифицированный интерфейс источников")
    p("Базовый класс EventSource с методом fetch() → list[RawEvent]. Каждый адаптер реализует своим способом.")
    b.h3("3.4.2. Адаптер RSS-лент")
    p("feedparser; для каждого URL из RSS_FEEDS получаем feed.entries, формируем RawEvent с title, summary, link, published.")
    b.h3("3.4.3. Адаптер Telegram-каналов")
    p(
        "Скачивание HTML страницы t.me/s/<channel> через httpx, BeautifulSoup "
        "извлекает сообщения. Это позволяет без api_id/api_hash работать с "
        "публичными каналами. На демо используются iteventsru и ITMeeting."
    )
    b.h3("3.4.4. Адаптеры платформ KudaGo и Lu.ma")
    p(
        "KudaGo — публичный JSON API, фильтр по категории «компьютеры». "
        "Lu.ma — ICS-фиды конкретных календарей: парсим как iCal, "
        "превращаем VEVENT в RawEvent. LUMA_CALENDARS — список URL'ов "
        "через запятую."
    )
    b.h3("3.4.5. AI-агент нормализации событий")
    p(
        "Одна нода LangGraph + один LLM-вызов. Входной prompt просит "
        "извлечь поля и оценить quality/hype в шкале 1–10. Выход обёрнут "
        "в pydantic.WithStructuredOutput — модель отдаёт уже типизированный "
        "объект. При сбое LLM (network, rate limit) автоматически "
        "переключается на groq_fallback_model."
    )

    b.h2("3.5. Реализация рекомендательной подсистемы")
    b.h3("3.5.1. Базовое ранжирование по правилам")
    p("Компонент rule в app/recommender/hybrid.py: совпадение тем (+2), формата (+3 или +1 для any), города (+2 или +1 для any). Возвращается с учётом score_rule_weight=0.5.")
    b.h3("3.5.2. Семантические представления и кэширование")
    p(
        "Embedding'и пользователя и событий кэшируются в таблицах. "
        "При запросе /recommendations выполняется два прогрева: "
        "refresh_user_embedding (с обязательным rollback'ом при сбое — это "
        "ключевая правка от 29 мая) и ensure_event_embeddings (пакетная "
        "досчёт недостающих)."
    )
    b.h3("3.5.3. Байесовская онлайн-адаптация на основе обратной связи")
    p(
        "user_topic_stats хранит alpha и beta по каждой теме. Лайк +3 к "
        "alpha, save +1, дизлайк +2 к beta. При чтении применяется "
        "temporal decay со скоростью gamma=0.995^days. При ранжировании — "
        "Thompson sampling: max_t Beta(alpha_t, beta_t)."
    )
    b.h3("3.5.4. Контекстный бандит LinUCB")
    p(
        "user_bandit_states хранит матрицу A (16×16) и вектор b (16). "
        "Контекстный вектор x формируется в context_vector: топ-3 темы, "
        "формат, город, freshness, hype, quality и др. На фидбэк: A += xxᵀ, "
        "b += reward·x. Скор — UCB: xᵀθ̂ + α·√(xᵀA⁻¹x), θ̂ = A⁻¹·b. "
        "Реализация — на чистом numpy."
    )
    b.h3("3.5.5. Коллаборативная компонента на основе LightGCN")
    p(
        "Тренируется отдельным скриптом scripts/train_gnn.py: строится "
        "двудольный граф user↔event по лайкам/saves, выполняется "
        "пропагация по графу за L=3 слоя, получаются эмбеддинги пользователей "
        "и событий. На малом датасете прототипа сигнал шумит — по умолчанию "
        "GNN_ENABLED=false. В режиме включения скор события считается как "
        "user_emb · event_emb."
    )
    b.h3("3.5.6. Гибридное ранжирование и MMR-диверсификация")
    p(
        "Итог = сумма девяти вкладов. Сортировка по убыванию. MMR-rerank "
        "(λ=0.7): каждый следующий элемент должен максимизировать "
        "λ·relevance − (1−λ)·max_similarity(уже_выбранные). Без MMR в "
        "топ-5 иногда попадали 5 митапов про один и тот же стек."
    )
    b.h3("3.5.7. Оптимизация производительности")
    p(
        "Главная оптимизация — устранение пересчёта эмбеддингов на каждом "
        "запросе. До оптимизации /recommendations занимал ~16 с на 200 "
        "событиях (encoder вызывался 200 раз). После — embedding'и "
        "кэшируются в БД, пользовательский считается один раз на запрос "
        "(а не на каждое событие), session_embedding для последних 5 "
        "взаимодействий blend'ится в пользовательский. Итог — ~0.4 с."
    )

    b.h2("3.6. Реализация мультиагентного Copilot")
    b.h3("3.6.1. Общая схема графа LangGraph")
    p(
        "StateGraph(CopilotState): START → retrieve → supervisor → "
        "conditional_edge → один из 5 специалистов → finalize → END. "
        "Узел retrieve собирает контекст (top-k событий по запросу, профиль, "
        "снимок взаимодействий). Supervisor классифицирует intent и "
        "возвращает имя следующего узла. Finalize возвращает {answer, "
        "recommended_event_ids, specialist}."
    )
    b.h3("3.6.2. Агент-супервизор и маршрутизация запросов")
    p(
        "Supervisor — отдельная нода с системным промптом «классифицируй "
        "запрос в один из intent'ов» и pydantic-схемой Intent. При сбое "
        "structured output (LLM вернула не валидный JSON) — keyword "
        "fallback по эвристикам. Conditional edge выбирает специалиста по "
        "supervisor.intent."
    )
    b.h3("3.6.3. Набор инструментов агентов")
    p(
        "Шесть function-calling tools в app/agents/copilot/tools.py: "
        "search_events(query, k), get_user_profile(telegram_id), "
        "get_interactions_summary(telegram_id), explain_event(event_id), "
        "recall_about_user(telegram_id, query, k), mark_saved(event_id). "
        "Связываются с LLM через bind_tools — "
        "модель сама решает, что вызвать."
    )

    b.h2("3.7. Реализация долговременной памяти пользователя")
    b.h3("3.7.1. Структура и хранение заметок")
    p(
        "Таблица user_memories: text, category (interest/dislike/goal/"
        "constraint/context/event_pref/other), salience (1–10), embedding, "
        "access_count, last_accessed_at, source, created_at. Создаются "
        "LLM-агентом из значимых фидбэков и реплик Copilot'а."
    )
    b.h3("3.7.2. Семантический поиск по памяти")
    p(
        "recall_about_user делает top-k по score = 0.7·cosine + "
        "0.3·salience_norm, обновляет access_count и last_accessed_at "
        "выбранных заметок (используется при compaction для приоритизации "
        "часто запрашиваемых)."
    )
    b.h3("3.7.3. Периодическое уплотнение памяти")
    p(
        "При достижении MEMORY_COMPACT_THRESHOLD (50 заметок) запускается "
        "compact_memories: LLM группирует заметки по категориям и сжимает "
        "до ≤5 на категорию (агрегируя факты). Compaction job в scheduler'е "
        "стоит на интервале раз в 7 дней; вручную — scripts/compact_memories.py."
    )

    b.h2("3.8. Реализация планировщика задач")
    b.h3("3.8.1. Архитектура процесса")
    p(
        "Отдельный процесс python -m app.scheduler.digest. Внутри — "
        "BlockingScheduler APScheduler с 5 джобами: daily_digest, "
        "ingest_habr, ingest_rss, ingest_telegram, compact_memories. "
        "Все три ingestion'а на демо запускаются сразу при старте "
        "(next_run_time=now) и потом повторяются по INGEST_INTERVAL_HOURS."
    )
    b.h3("3.8.2. Рассылка персонализированного дайджеста")
    p(
        "send_digest_once: достать список подписчиков (/subscriptions/users), "
        "для каждого дёрнуть /agent-recommendations/{tid}/cards, "
        "взять топ-1 карточку, отправить через Telegram Bot API. Интервал "
        "и запуск-на-старте конфигурируются через DIGEST_INTERVAL_MINUTES "
        "и DIGEST_RUN_ON_STARTUP."
    )
    b.h3("3.8.3. Периодический ингест и уплотнение памяти")
    p(
        "ingest_habr_once, ingest_rss_once, ingest_telegram_once дёргают "
        "соответствующие /ingestion/... POST-endpoint'ы API. compact_memories "
        "job — раз в 7 дней — вызывает app.recommender.memory.compact_*."
    )

    b.h2("3.9. Поток данных в системе")
    p(
        "Источник → raw_event (status=raw) → AI-нормализация → events / "
        "raw_event=non_it. Пользователь жмёт «🎯 Рекомендации» → бот "
        "→ API → hybrid_score по 9 компонент → MMR → JSON → бот рисует "
        "карточку. Пользователь жмёт 👍 → бот → API → 4 подсистемы "
        "обновляются → возврат. /copilot → бот → API → LangGraph → "
        "specialist → tools → finalize → бот."
    )

    # ════════════════════════════════════════════════════════════════════
    # Глава 4
    # ════════════════════════════════════════════════════════════════════
    b.h1("Глава 4. Тестирование")
    b.h2("4.1. Функциональное тестирование")
    b.h3("4.1.1. Организация тестового набора")
    p("180 тестов в 25 файлах. Каждый файл — отдельный модуль/подсистема. Прогон через pytest -q занимает ~30 с.")
    b.h3("4.1.2. Тесты алгоритмов рекомендательной подсистемы")
    p(
        "test_scoring (правила и hybrid base), test_bayesian + test_bayes_decay "
        "(α/β + temporal decay), test_bandit (LinUCB inference/update), "
        "test_gnn (LightGCN forward), test_multi_objective (взвешенная "
        "сумма), test_cold_start (/bio + skill-профиль)."
    )
    b.h3("4.1.3. Тесты подсистемы ингеста и нормализации")
    p("test_ingestion (HTTP-моки источников), test_multi_source (комбинирование), test_dedup (cosine ≥ 0.92).")
    b.h3("4.1.4. Тесты мультиагентного Copilot")
    p("test_supervisor (intent-классификация), test_specialists (отдельные ноды), test_copilot_tools (tools), test_copilot_graph_e2e (граф целиком).")

    b.h2("4.2. Офлайн-оценка рекомендательной системы")
    b.h3("4.2.1. Методика leave-one-out")
    p("Для каждой interaction типа like/save: убрать её, переранжировать, посмотреть, попадает ли событие в top-k.")
    b.h3("4.2.2. Сравнение алгоритмов по метрикам Recall@k и nDCG@k")
    p("Сравниваются: rule, content-only, hybrid_no_bandit, hybrid_full. Hybrid_full даёт лучшие Recall@5 и nDCG@5.")
    b.h3("4.2.3. Оценка качества с помощью LLM-судьи")
    p(
        "scripts/llm_judge.py берёт топ-N выдачи и просит LLM оценить "
        "«осмысленность ленты» с указанием конкретных проблем. Запускается "
        "вручную; результат — JSON-отчёт."
    )

    b.h2("4.3. Пользовательское тестирование")
    b.h3("4.3.1. Методика")
    p("Прогон сценариев из главы 1.7 на реальном Telegram-боте с реальным аккаунтом автора и тестовым набором событий.")
    b.h3("4.3.2. Тестируемые сценарии")
    li("Полный onboarding с настройкой профиля")
    li("Получение ленты и обратная связь (👍/👎/⭐)")
    li("Команда /undo")
    li("Поиск keyword + семантический")
    li("Cold-start через /bio")
    li("Мульти-туровый диалог с Copilot")
    li("Подписка и получение дайджеста")
    b.h3("4.3.3. Выявленные дефекты и их устранение")
    p(
        "Главные обнаруженные на майской итерации: 1) generic-сообщение "
        "«Copilot временно недоступен» без логов (фикс — logger.exception "
        "во всех ветках); 2) дубль кнопки «Начать настройку» (фикс — "
        "ReplyKeyboardRemove на /start); 3) пропадание кнопки запуска "
        "настройки при пропуске тура (фикс — отдельное сообщение с "
        "инлайн-кнопкой); 4) пустая лента при настроенном профиле — "
        "PendingRollbackError из-за непровёрнутого flush в "
        "refresh_user_embedding (фикс — обязательный rollback на сбое); "
        "5) англицизмы в «Почему» (фикс — перепись подписей в "
        "_COMPONENT_LABELS_RU); 6) DIGEST_INTERVAL_MINUTES для dev-режима."
    )

    b.h2("4.4. Вывод")
    p(
        "Все функциональные требования реализованы и покрыты тестами. "
        "Hybrid_full демонстрирует лучшие офлайн-метрики среди вариантов "
        "конфигурации. Пользовательское тестирование выявило шесть "
        "дефектов; все исправлены в майской итерации 2026."
    )

    # ════════════════════════════════════════════════════════════════════
    # Глава 5
    # ════════════════════════════════════════════════════════════════════
    b.h1("Глава 5. Доработки после первой версии прототипа")
    b.h2("5.1. Устранение дефектов и рефакторинг рекомендера")
    p(
        "Главное — устранение «тихих» сбоев. В refresh_user_embedding "
        "теперь обязательный rollback при «database is locked»; "
        "это убирает PendingRollbackError, который ранее приводил к 500 "
        "на /recommendations при настроенном профиле. В /copilot и "
        "/agent-recommendations добавлены logger.exception во всех ветках "
        "fallback'а — раньше generic-сообщение приходило в UI без "
        "записи в логи."
    )
    b.h2("5.2. Поддержка PostgreSQL + pgvector")
    p(
        "Миграция 9a1b2c3d4e5f добавляет тип vector(384) для embedding-"
        "колонок в PG. pgvector_top_k_events использует оператор <=> для "
        "быстрого top-k на стороне БД. На SQLite тип остаётся TEXT и "
        "сортировка — in-process."
    )
    b.h2("5.3. Health-эндпоинт и fallback LLM")
    p("/health возвращает {db, llm, ingestion: ok/fail}. LLM — однократный fallback с primary на groq_fallback_model.")
    b.h2("5.4. Ускорение AI-рекомендаций (вариант C)")
    p(
        "Старая 4-нодовая цепочка LangGraph (user_profile → event_analyzer "
        "→ recommendation → explanation) заменена на: hybrid отбирает top-N, "
        "один LLM-вызов пишет объяснения батчем. Параметр limit реально "
        "влияет."
    )
    b.h2("5.5. Расширение пользовательского интерфейса")
    p(
        "Мульти-туровый Copilot (handlers/copilot.py, активная сессия 15 "
        "мин, продолжение простыми сообщениями). Карточка события: вместо "
        "переключения mode=rule/ai — два независимых пояснения (❓ Почему "
        "+ 📖 Подробнее). Кнопка «🚀 Начать настройку» появляется даже "
        "при пропуске тура. На /start сбрасывается старая reply-клавиатура."
    )
    b.h2("5.6. Качество кода")
    p("ruff — без ошибок на app/. 285 тестов проходят. CI: два job'а — test-sqlite (быстрый default) и test-postgres (на образе pgvector/pgvector:pg16) — критичные backend-зависимые пути проверяются на реальном PG. Все хендлеры покрыты pytest-фикстурами с in-memory SQLite.")
    b.h2("5.7. Итог")
    p(
        "Майская итерация привела систему к стабильному состоянию: "
        "критические пути защищены fallback'ами, объяснимость реализована "
        "на двух уровнях, копилот мульти-туровый, dev-окружение настраивается "
        "через .env (включая 10-минутный дайджест), документация (этот "
        "отчёт, обзор и план показа) приведена в соответствие с реализацией."
    )

    # ── Заключение ───────────────────────────────────────────────────────
    b.h1("Заключение")
    p(
        "В рамках курсовой работы спроектирована и реализована программная "
        "система EventMind для агрегации IT-мероприятий и персонализированных "
        "AI-рекомендаций. Решены все поставленные задачи: проведён анализ "
        "предметной области и сформированы требования; спроектирована "
        "архитектура из трёх независимых процессов и схема БД из 11 таблиц; "
        "реализован пайплайн ingestion с шестью адаптерами и AI-нормализатором; "
        "построен гибридный рекомендатель из девяти компонент, объединяющий "
        "content-based, контекстный бандит (LinUCB), коллаборативный "
        "фильтр (LightGCN), Bayesian preference (Thompson sampling) и "
        "эвристики качества; реализован мульти-агентный AI-копилот на "
        "LangGraph; обеспечены REST API, Telegram-бот и контейнеризованное "
        "развёртывание."
    )
    p(
        "Научно-практическая новизна работы — одновременная реализация "
        "контекстной и коллаборативной парадигм в одном ранжировании, "
        "двухуровневая объяснимость выдачи (короткий поп-ап и развёрнутый "
        "разбор с counterfactual) и долговременная mem0-стилевая память "
        "пользователя, подключённая к AI-копилоту. Все компоненты покрыты "
        "автотестами (180 кейсов) и офлайн-оценкой (Recall@k, nDCG@k)."
    )
    p(
        "Дальнейшее развитие: 1) накопление достаточного объёма данных "
        "для эффективной работы GNN-компоненты; 2) расширение набора "
        "ingestion-источников за пределы шести текущих; 3) автоматическая "
        "оптимизация весов hybrid через bandit поверх самих компонент; 4) "
        "доработка LLM-судьи в эвале до периодического pipeline."
    )

    # ── Глоссарий ────────────────────────────────────────────────────────
    b.h1("Глоссарий")
    li("LLM — большая языковая модель; в проекте цепочка _LLMChain: Gemini (REST + autopobe) → Groq llama-3.3-70b → Groq llama-3.1-8b")
    li("Embedding — векторное представление текста (384 dim, MiniLM-L12-v2)")
    li("Cosine similarity — мера близости двух векторов")
    li("Hybrid recommender — комбинация разных алгоритмов в один score (у нас 9 компонент)")
    li("MMR — Maximal Marginal Relevance, пост-процессинг для разнообразия")
    li("Thompson sampling — Bayesian explore/exploit стратегия")
    li("LinUCB — контекстный бандит с UCB-бонусом за неопределённость")
    li("LightGCN — графовая коллаборативная модель")
    li("RAG — Retrieval-Augmented Generation, retrieve → подмешать в промпт → generate")
    li("Function calling — LLM сама решает, какую функцию вызвать")
    li("LangGraph — фреймворк для multi-agent графов")
    li("Supervisor-Worker — паттерн multi-agent: координатор + специалисты")
    li("Cold start — холодный старт нового пользователя без истории")
    li("Long-term memory (mem0) — активные заметки про пользователя от первого лица")
    li("pgvector — расширение PostgreSQL для типа vector + оператора <=> (cosine distance)")
    li("APScheduler — Python-планировщик фоновых задач по интервалу/cron")

    # ── Библиографический список ─────────────────────────────────────────
    b.h1("Библиографический список")
    li("Aggarwal C. C. Recommender Systems: The Textbook. — Springer, 2016.")
    li("Ricci F., Rokach L., Shapira B. Recommender Systems Handbook, 3rd Ed. — Springer, 2022.")
    li("Burke R. Hybrid Recommender Systems: Survey and Experiments. — UMUAI 12(4), 2002.")
    li("Li L., Chu W., Langford J., Schapire R. A Contextual-Bandit Approach to Personalized News Article Recommendation. — WWW 2010.")
    li("Carbonell J., Goldstein J. The Use of MMR, Diversity-Based Reranking for Reordering Documents. — SIGIR 1998.")
    li("Reimers N., Gurevych I. Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks. — EMNLP 2019.")
    li("He X. et al. LightGCN: Simplifying and Powering Graph Convolution Network for Recommendation. — SIGIR 2020.")
    li("Chen S. mem0 — long-term memory layer for LLM apps. — arXiv 2024.")
    li("LangGraph documentation. — https://langchain-ai.github.io/langgraph/")
    li("FastAPI documentation. — https://fastapi.tiangolo.com/")
    li("aiogram 3 documentation. — https://docs.aiogram.dev/")
    li("SQLAlchemy 2.0 documentation. — https://docs.sqlalchemy.org/")
    li("pgvector. — https://github.com/pgvector/pgvector")


# ════════════════════════════════════════════════════════════════════════
#  Main
# ════════════════════════════════════════════════════════════════════════


def main() -> None:
    # 1. Обзор проекта — стили H1/H2/H3 + Normal/List Bullet
    overview_path = DOCS / "Обзор_проекта_EventMind.docx"
    doc1 = Document(overview_path)
    _clear_body(doc1)
    b1 = Builder(doc1, h1="H1", h2="H2", h3="H3", normal="Normal", bullet="List Bullet")
    build_overview(b1)
    doc1.save(overview_path)
    print(f"updated: {overview_path}")

    # 2. План показа — стили Heading 1/2/3 + Normal/List Bullet
    plan_path = DOCS / "План_показа_EventMind.docx"
    doc2 = Document(plan_path)
    _clear_body(doc2)
    b2 = Builder(doc2, h1="Heading 1", h2="Heading 2", h3="Heading 3", normal="Normal", bullet="List Bullet")
    build_presentation_plan(b2)
    doc2.save(plan_path)
    print(f"updated: {plan_path}")

    # 3. Отчёт — стили HSE Title 1/2/3 + HSE Default Text + HSE list
    thesis_path = DOCS / "отчет_Курсовая.docx"
    doc3 = Document(thesis_path)
    _clear_body(doc3)
    b3 = Builder(
        doc3,
        h1="HSE Title 1", h2="HSE Title 2", h3="HSE Title 3",
        normal="HSE Default Text", bullet="HSE list",
    )
    build_thesis(b3, default_text="HSE Default Text", list_style="HSE list")
    doc3.save(thesis_path)
    print(f"updated: {thesis_path}")


if __name__ == "__main__":
    main()
